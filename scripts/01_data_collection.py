import os
import random
import requests
import zipfile
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
import json

# 创建数据目录
os.makedirs("data/raw_resumes", exist_ok=True)
os.makedirs("data/raw_jobs", exist_ok=True)
os.makedirs("data/processed", exist_ok=True)

def download_kaggle_dataset(dataset_slug, download_path="data"):
    """
    从Kaggle下载数据集。
    需要您在系统中配置Kaggle API凭证(kaggle.json)。
    """
    print(f"正在从Kaggle下载数据集: {dataset_slug}")
    try:
        # 使用Kaggle API下载数据集
        # 注意：这需要在运行环境正确配置Kaggle API凭证
        os.system(f"kaggle datasets download -d {dataset_slug} -p {download_path} --unzip")
        print("数据集下载并解压成功。")
    except Exception as e:
        print(f"Kaggle数据集下载失败: {e}")
        print("请确保您已安装Kaggle API (pip install kaggle)并已配置API凭证。")
        print("您可以从Kaggle账户页面创建API令牌并将其放在~/.kaggle/kaggle.json。")

def generate_mock_resumes(num_resumes=200, seed=42):
    """生成多样化模拟简历数据（docx），并同时构造可用于实体训练的文本与标注。"""
    random.seed(seed)
    print(f"正在生成 {num_resumes} 份多样化简历...")
    names = ["张三", "李四", "王五", "赵六", "钱七", "孙八", "周九", "吴十", "郑一", "冯二", "陈三", "楮四", "卫五"]
    positions = [
        "后端工程师", "前端工程师", "算法工程师", "数据工程师", "数据科学家",
        "机器学习工程师", "NLP工程师", "计算机视觉工程师", "DevOps工程师", "架构师"
    ]
    degrees = ["博士", "硕士", "本科", "大专"]
    majors = ["计算机科学", "软件工程", "人工智能", "信息工程", "数学", "统计学", "电子工程"]
    schools = ["清华大学", "北京大学", "上海交通大学", "浙江大学", "复旦大学", "中国科学技术大学", "南京大学"]
    companies = ["字节跳动", "阿里巴巴", "腾讯", "华为", "美团", "百度", "京东", "小米"]
    skill_pool = [
        "Python", "Java", "Go", "C++", "C#", "Rust", "Scala",
        "TensorFlow", "PyTorch", "scikit-learn", "XGBoost", "LightGBM",
        "Docker", "Kubernetes", "Linux", "Git", "MySQL", "PostgreSQL", "Redis", "MongoDB",
        "Hadoop", "Spark", "Flink", "Kafka", "Airflow",
        "Vue", "React", "Angular", "TypeScript", "Node.js",
        "Elasticsearch", "Faiss", "Milvus", "LangChain", "RAG",
        "AWS", "Azure", "GCP", "Terraform", "Jenkins"
    ]
    languages = ["英语四级", "英语六级", "CET-6", "雅思", "托福", "日语N1"]
    certs = ["PMP", "CKA", "CKAD", "AWS SAA", "RHCE"]

    # 同时构造用于实体训练的样本
    train_samples = []

    for i in range(num_resumes):
        name = f"{random.choice(names)}_{i}"
        position = random.choice(positions)
        degree = random.choice(degrees)
        major = random.choice(majors)
        years = random.randint(0, 12)
        picked_skills = ", ".join(random.sample(skill_pool, k=random.randint(3, 6)))
        company = random.choice(companies)
        school = random.choice(schools)
        salary = f"{random.randint(15, 45)}K-{random.randint(46, 80)}K"
        lang = random.choice(languages)
        cert = random.choice(certs) if random.random() < 0.4 else ""

        # 构造文本并记录实体边界
        lines = []
        entities = []
        offset = 0

        def add_line(label, value, ent_type=None):
            nonlocal offset
            line = f"{label}：{value}\n"
            # 记录实体范围（仅在提供 ent_type 时）
            if ent_type and value:
                start = offset + len(label) + 1  # “label：” 的长度
                end = start + len(value)
                entities.append({"start": start, "end": end, "type": ent_type})
            lines.append(line)
            offset += len(line)

        add_line("姓名", name, "姓名")
        add_line("求职意向", position, "职位")
        add_line("学历", degree, "学历")
        add_line("专业", major, "专业")
        add_line("工作年限", f"{years}年", "工作年限")
        add_line("技能", picked_skills, "技能")
        add_line("项目经验", f"在{company}参与{position}相关项目，负责模块开发与优化", "项目经验")
        add_line("公司名称", company, "公司名称")
        add_line("毕业院校", school, "毕业院校")
        add_line("薪资期望", salary, "薪资期望")
        if cert:
            add_line("证书", cert, None)  # 证书不在当前实体集合中，但保留文本
        add_line("语言能力", lang, None)

        full_text = "".join(lines)

        # 保存 docx 文件
        doc = Document()
        doc.add_heading("个人简历", 0)
        for ln in full_text.split("\n"):
            if ln:
                doc.add_paragraph(ln)
        doc.save(f"data/raw_resumes/mock_resume_{i}.docx")

        # 收集训练样本（用于 entity_train.json）
        train_samples.append({"text": full_text, "entities": entities})

    # 输出训练集
    out_path = os.path.join("data/processed", "entity_train.json")
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(train_samples, f, ensure_ascii=False, indent=2)
    print(f"模拟简历生成完毕，并已导出实体训练数据: {out_path}")


def generate_mock_jobs(num_jobs=30, seed=7):
    """生成多样化岗位JD（txt）。"""
    random.seed(seed)
    print(f"正在生成 {num_jobs} 份岗位JD...")
    roles = [
        {
            "name": "backend_engineer",
            "title": "后端工程师",
            "skills": ["Java", "Go", "MySQL", "Redis", "Kafka"],
            "keywords": ["微服务", "REST", "高并发"],
            "years": (2, 5),
            "degree": "本科",
            "extras": ["了解 Docker/Kubernetes，具备云原生经验。"]
        },
        {
            "name": "frontend_engineer",
            "title": "前端工程师",
            "skills": ["TypeScript", "React", "Webpack", "Node.js"],
            "keywords": ["组件化", "性能优化"],
            "years": (1, 4),
            "degree": "本科",
            "extras": ["有跨端或SSR经验者优先。"]
        },
        {
            "name": "data_engineer",
            "title": "数据工程师",
            "skills": ["Spark", "Hadoop", "Kafka", "Airflow"],
            "keywords": ["数据仓库", "ETL"],
            "years": (3, 6),
            "degree": "本科",
            "extras": ["熟悉Linux环境与脚本，掌握SQL优化。"]
        },
        {
            "name": "ml_engineer",
            "title": "机器学习工程师",
            "skills": ["Python", "PyTorch", "TensorFlow", "scikit-learn"],
            "keywords": ["推荐系统", "NLP", "LLM"],
            "years": (2, 5),
            "degree": "硕士",
            "extras": ["掌握RAG/向量检索，熟悉LangChain优先。"]
        },
        {
            "name": "devops_engineer",
            "title": "DevOps工程师",
            "skills": ["Docker", "Kubernetes", "Jenkins", "Terraform"],
            "keywords": ["CI/CD", "云原生"],
            "years": (2, 6),
            "degree": "本科",
            "extras": ["具备云平台经验（AWS/Azure/GCP）。"]
        }
    ]

    for i in range(num_jobs):
        role = random.choice(roles)
        yrs = random.randint(role["years"][0], role["years"][1])
        ks = ", ".join(role["skills"]) 
        kws = ", ".join(role["keywords"]) 
        lines = [
            f"岗位：{role['title']}",
            f"学历要求：{role['degree']}",
            f"工作年限：{yrs}年",
            f"技能要求：{ks}",
            f"方向关键词：{kws}",
        ]
        lines.extend(role["extras"]) 
        content = "\n".join(lines)
        with open(f"data/raw_jobs/mock_job_{role['name']}_{i}.txt", "w", encoding="utf-8") as f:
            f.write(content)
    print("岗位JD生成完毕。")

def scrape_job_postings(job_title="算法工程师", num_pages=1):
    """
    （示例）从招聘网站爬取岗位描述。
    注意：爬取网站数据需遵守网站的robots.txt协议，并可能需要处理反爬机制。
    此函数仅为基本示例，实际使用时需要进行适配和扩展。
    """
    print(f"正在爬取'{job_title}'的岗位信息...")
    # 以某招聘网站为例
    base_url = "https://search.51job.com/list/000000,000000,0000,00,9,99,{job_title},2,{page}.html"

    for page in range(1, num_pages + 1):
        url = base_url.format(job_title=job_title, page=page)
        try:
            response = requests.get(url)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, "html.parser")
            
            # 查找岗位链接
            # 注意：这里的选择器需要根据目标网站的实际HTML结构进行调整
            job_links = soup.select(".el .t a") 

            for link in job_links:
                job_url = link['href']
                # 爬取岗位详情页
                # ... 此处省略详情页爬取逻辑 ...
                # 将爬取到的岗位描述保存到 data/raw_jobs 目录
                
        except requests.exceptions.RequestException as e:
            print(f"请求失败: {e}")
        except Exception as e:
            print(f"爬取过程中出现错误: {e}")
    print("岗位信息爬取完成（示例）。")


if __name__ == "__main__":
    # 1.（可选）下载公开数据集
    # download_kaggle_dataset("jillanisofttech/resume-dataset")

    # 2. 生成多样化简历与实体训练数据
    generate_mock_resumes(num_resumes=200)

    # 3. 生成多样化岗位JD
    generate_mock_jobs(num_jobs=30)

    # 4.（可选）爬取岗位数据（默认不执行；请遵守站点策略）
    # scrape_job_postings(job_title="算法工程师", num_pages=1)
    # scrape_job_postings(job_title="产品经理", num_pages=1)
    # scrape_job_postings(job_title="运营专员", num_pages=1)