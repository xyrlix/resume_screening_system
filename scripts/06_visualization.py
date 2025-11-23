import os
import json
import time
import math
import streamlit as st
import urllib.request
import urllib.error
import importlib.util
from typing import Tuple, List, Dict
import matplotlib
import matplotlib.pyplot as plt
import matplotlib.font_manager

# 设置matplotlib支持中文字体
# 检查系统中可用的中文字体
chinese_fonts = ['SimHei', 'Microsoft YaHei', 'STHeiti', 'Arial Unicode MS', 'DejaVu Sans']
available_fonts = []
for font in chinese_fonts:
    try:
        matplotlib.font_manager.findfont(matplotlib.font_manager.FontProperties(family=font))
        available_fonts.append(font)
    except:
        pass

# 设置可用的中文字体
if available_fonts:
    matplotlib.rcParams['font.sans-serif'] = available_fonts + ['sans-serif']
matplotlib.rcParams['axes.unicode_minus'] = False

try:
    from streamlit_echarts import st_echarts
except Exception:
    st_echarts = None


def page_setup():
    st.set_page_config(page_title="智能简历筛选系统", page_icon="📄", layout="wide")
    st.title("📄 智能简历筛选系统")
    st.caption("简历与岗位匹配，可视化展示匹配结果")
    st.markdown(
        """
        <style>
        .stTabs [role="tab"] {
            font-size: 18px;
            font-weight: 600;
            color: #1976d2;
        }
        .stTabs [role="tab"][aria-selected="true"] {
            color: #1976d2;
            border-bottom: 3px solid #1976d2;
        }
        .step-title {
            padding: 6px 10px;
            border-left: 4px solid #1976d2;
            font-weight: 700;
            font-size: 18px;
        }
        h3 {
            color: #1976d2;
            font-weight: 700;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def try_health(url: str, timeout: float = 0.8) -> bool:
    try:
        req = urllib.request.Request(url)
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return resp.status == 200
    except Exception:
        return False


def detect_api_base() -> str:
    # 优先使用环境变量（例如在部署时配置）
    env_base = os.getenv("API_BASE_URL")
    if env_base and try_health(env_base.rstrip("/") + "/health"):
        return env_base.rstrip("/")
    # 尝试本地端口范围（与后端自动回退一致）
    for port in range(8000, 8006):
        base = f"http://127.0.0.1:{port}"
        if try_health(base + "/health"):
            return base
    # 兜底：默认 8000
    return "http://127.0.0.1:8000"


def api_post(base: str, path: str, payload: dict, timeout: float | None = None):
    url = base.rstrip("/") + path
    data = json.dumps(payload).encode("utf-8")
    headers = {"Content-Type": "application/json"}
    tok = st.session_state.get("api_token")
    if tok:
        headers["Authorization"] = f"Bearer {tok}"
    req = urllib.request.Request(url, data=data, headers=headers)
    to = timeout if timeout is not None else float(st.session_state.get("api_timeout", 2.5))
    try:
        with urllib.request.urlopen(req, timeout=to) as resp:
            return json.loads(resp.read().decode("utf-8")), None
    except urllib.error.HTTPError as e:
        return None, f"HTTPError {e.code}: {e.reason}"
    except Exception as e:
        return None, str(e)

def api_get(base: str, path: str, timeout: float | None = None):
    url = base.rstrip("/") + path
    headers = {"Accept": "application/json"}
    tok = st.session_state.get("api_token")
    if tok:
        headers["Authorization"] = f"Bearer {tok}"
    req = urllib.request.Request(url, headers=headers)
    to = timeout if timeout is not None else float(st.session_state.get("api_timeout", 2.5))
    try:
        with urllib.request.urlopen(req, timeout=to) as resp:
            return json.loads(resp.read().decode("utf-8")), None
    except urllib.error.HTTPError as e:
        return None, f"HTTPError {e.code}: {e.reason}"
    except Exception as e:
        return None, str(e)

def _download_excel_or_csv_sheets(sheets: Dict[str, List[Dict]], excel_name: str, csv_prefix: str, prefer_excel: bool = True):
    import pandas as pd, io, importlib.util
    try:
        if not prefer_excel or importlib.util.find_spec("xlsxwriter") is None:
            raise ImportError("xlsxwriter not available")
        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine="xlsxwriter") as w:
            for name, rows in sheets.items():
                pd.DataFrame(rows).to_excel(w, index=False, sheet_name=name)
        st.download_button("下载Excel", data=buf.getvalue(), file_name=excel_name)
    except Exception:
        for name, rows in sheets.items():
            import pandas as pd
            df = pd.DataFrame(rows)
            csv = df.to_csv(index=False).encode("utf-8-sig")
            st.download_button(f"下载{name} CSV", data=csv, file_name=f"{csv_prefix}_{name}.csv")

def _post_with_progress(base: str, path: str, payload: dict, timeout: float, label: str):
    import concurrent.futures, time
    bar = st.progress(0.0)
    to = max(float(timeout or 10.0), 1.0)
    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
        fut = ex.submit(api_post, base, path, payload, to)
        start = time.time()
        while not fut.done():
            elapsed = time.time() - start
            pct = min(0.95, elapsed / to)
            bar.progress(pct)
            time.sleep(0.1)
        res, err = fut.result()
        bar.progress(1.0)
        return res, err


def _load_module(module_filename: str):
    base_dir = os.path.dirname(__file__)
    path = os.path.join(base_dir, module_filename)
    spec = importlib.util.spec_from_file_location(module_filename.replace('.py',''), path)
    mod = importlib.util.module_from_spec(spec)
    if spec and spec.loader:
        spec.loader.exec_module(mod)
    return mod


# 预加载用于本地匹配与文件解析的模块
_matcher_mod = None
_preprocess_mod = None
try:
    _matcher_mod = _load_module('04_matcher_model.py')
except Exception as e:
    _matcher_mod = None
try:
    _preprocess_mod = _load_module('02_data_preprocess.py')
except Exception:
    _preprocess_mod = None


def sidebar():
    if "api_base" not in st.session_state:
        st.session_state.api_base = detect_api_base()
    with st.sidebar:
        st.subheader("身份与角色")
        if "role" not in st.session_state:
            st.session_state["role"] = "面向招聘方"
        role_options = ["面向招聘方", "面向求职者", "管理员"]
        sel = st.radio("选择角色", role_options, index=role_options.index(st.session_state["role"]) if st.session_state["role"] in role_options else 0, key="role_radio")
        st.session_state["role"] = sel
        st.markdown(f"<div style='margin-top:6px;padding:8px 10px;border-radius:6px;background:#1976d2;color:#fff;font-weight:600;display:inline-block;'>当前角色：{st.session_state['role']}</div>", unsafe_allow_html=True)
        st.markdown("---")
        st.subheader("⚙️ 服务与配置")
        c0, c1 = st.columns([3, 1])
        with c0:
            api_base = st.text_input("API地址", st.session_state.api_base)
        with c1:
            refresh = st.button("检测", key="btn_health_check", type="secondary")
        base = (api_base or st.session_state.api_base).rstrip("/")
        ok = try_health(base + "/health")
        st.metric("API健康", "正常" if ok else "异常")
        st.link_button("查看API文档", base + "/docs")
        cfg_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "config", "matching.json")
        if os.path.isfile(cfg_path):
            with open(cfg_path, "r", encoding="utf-8") as f:
                try:
                    cfg = json.load(f)
                    st.expander("匹配配置预览").json(cfg)
                except Exception:
                    st.info("配置读取失败")
        if refresh:
            if try_health(base + "/health"):
                st.success("API健康正常")
                st.session_state.api_base = base
            else:
                auto = detect_api_base()
                if try_health(auto + "/health"):
                    st.session_state.api_base = auto
                    st.success(f"已自动切换到 {auto}")
                else:
                    st.error("API不可用，请检查后端服务或端口")
        st.caption("提示：可通过 API_HOST/API_PORT 控制后端监听；未设置时会自动在 8000–8005 中选择可用端口。")
        st.markdown("---")
        st.subheader("全局设置")
        api_to = st.slider("默认API超时(秒)", 2.0, 30.0, float(st.session_state.get("api_timeout", 12.0)), 0.5, key="sidebar_api_timeout")
        st.session_state["api_timeout"] = float(api_to)
        llm_enabled = st.checkbox("启用大模型(全局)", value=bool(st.session_state.get("llm_global_enabled", False)), key="sidebar_llm_enabled")
        adv = st.expander("高级设置", expanded=False)
        with adv:
            api_token = st.text_input("API令牌 (可选)", value=st.session_state.get("api_token", ""), help="用于访问需要鉴权的接口，作为 Authorization: Bearer <token> 传递")
            export_csv_default = st.checkbox("默认导出为CSV", value=bool(st.session_state.get("export_csv_default", False)), key="sidebar_export_csv")
            llm_provider = st.selectbox("LLM Provider", options=["local", "qwen", "openai", "ark", "hunyuan", "qianfan", "kimi", "gemini", "tavily"], index=["local", "qwen", "openai", "ark", "hunyuan", "qianfan", "kimi", "gemini", "tavily"].index(st.session_state.get("llm_provider", "local")), key="llm_provider_sel")
            llm_model = st.text_input("LLM模型名", value=st.session_state.get("llm_model", ""), key="llm_model_name")
            llm_api_url = st.text_input("LLM API URL", value=st.session_state.get("llm_api_url", ""), key="llm_api_url")
            llm_api_key = st.text_input("LLM API KEY", value=st.session_state.get("llm_api_key", ""), key="llm_api_key")
            prov_status = "云端就绪" if llm_provider != "local" and (llm_api_key or os.getenv(f"{llm_provider.upper()}_API_KEY", "")) else "本地回退"
            st.metric("LLM Provider", f"{llm_provider} - {prov_status}")
        apply = st.button("应用全局设置", key="sidebar_apply", type="secondary")
        if apply:
            st.session_state["export_csv_default"] = bool(export_csv_default)
            st.session_state["llm_global_enabled"] = bool(llm_enabled)
            if "api_token" in locals():
                st.session_state["api_token"] = api_token
            st.session_state["llm_provider"] = llm_provider
            st.session_state["llm_model"] = llm_model
            _ = api_post(st.session_state.api_base, "/config/llm_enabled", {"enabled": bool(llm_enabled)})
            _ = api_post(st.session_state.api_base, "/config/llm_settings", {"provider": llm_provider, "api_url": llm_api_url, "api_key": llm_api_key, "model": llm_model})


def page_health():
    st.subheader("服务概览")
    base = st.session_state.api_base
    col1, col2 = st.columns(2)
    with col1:
        st.write("后端 API:", base)
        ok = try_health(base + "/health")
        st.metric("API健康", "正常" if ok else "异常")
        st.link_button("查看API文档", base + "/docs")
    with col2:
        st.write("匹配配置文件:", "config/matching.json")
        cfg_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "config", "matching.json")
        if os.path.isfile(cfg_path):
            with open(cfg_path, "r", encoding="utf-8") as f:
                cfg = json.load(f)
            st.json(cfg)
        else:
            st.info("未找到匹配配置文件，使用脚本内置默认配置。")
    st.markdown("---")
    st.subheader("预热向量索引")
    resume_json = os.path.join("data", "processed", "resumes_for_annotation.json")
    run = st.button("从处理后的简历入库向量", key="warmup_vector")
    if run:
        if not os.path.isfile(resume_json):
            st.error("简历JSON不存在")
            return
        with open(resume_json, "r", encoding="utf-8") as f:
            data = json.load(f)
        items = [{"id": str(i), "text": d.get("text", "")} for i, d in enumerate(data)]
        res, err = api_post(base, "/vector_index", {"items": items}, timeout=12.0)
        if err:
            st.error(f"入库失败：{err}")
        else:
            st.success(f"入库完成：{res.get('count')} 条")


def page_predict():
    st.subheader("实体预测（NER）")
    base = st.session_state.api_base
    text = st.text_area("输入简历文本", height=180, placeholder="在此粘贴简历文本…")
    if st.button("调用 /predict"):
        if not text.strip():
            st.warning("请先输入简历文本")
        else:
            res, err = api_post(base, "/predict", {"text": text})
            if err:
                st.error(f"调用失败：{err}")
            else:
                ents = res.get("entities", [])
                st.success(f"提取到 {len(ents)} 个实体")
                if ents:
                    st.dataframe(ents, use_container_width=True)


def page_match_single():
    st.subheader("单次匹配（Resume vs Job）")
    base = st.session_state.api_base
    c1, c2 = st.columns(2)
    with c1:
        resume_text = st.text_area("简历文本", height=200, placeholder="在此粘贴简历文本…")
        rfile = st.file_uploader("上传简历文件（txt/md/docx/pdf）", type=["txt", "md", "docx", "pdf"], key="resume_file")
        if rfile is not None:
            resume_text = _read_uploaded_text(rfile)
            st.info("已从上传文件填充简历文本")
    with c2:
        job_text = st.text_area("岗位文本", height=200, placeholder="在此粘贴岗位描述…")
        jfile = st.file_uploader("上传岗位文件（txt/md/docx/pdf）", type=["txt", "md", "docx", "pdf"], key="job_file")
        if jfile is not None:
            job_text = _read_uploaded_text(jfile)
            st.info("已从上传文件填充岗位文本")

    use_temp_weights = st.checkbox("使用临时权重在本地计算（不调用后端）", value=False)
    if st.button("调用 /match"):
        if not resume_text.strip() or not job_text.strip():
            st.warning("请同时输入简历与岗位文本")
        else:
            if use_temp_weights and _matcher_mod is not None and "temp_weights" in st.session_state:
                res = _local_match_with_weights(resume_text, job_text, st.session_state["temp_weights"])
                _render_match_result(res)
            else:
                res, err = _post_with_progress(base, "/match", {"resume_text": resume_text, "job_text": job_text}, float(st.session_state.get("api_timeout", 12.0)), "匹配评估")
                if err:
                    st.error(f"调用失败：{err}")
                else:
                    _render_match_result(res)


def page_match_batch():
    st.subheader("批量匹配（目录）")
    base = st.session_state.api_base
    default_resume_json = os.path.join("data", "processed", "resumes_for_annotation.json")
    default_jobs_dir = os.path.join("data", "raw_jobs")
    resume_json = st.text_input("简历JSON路径", default_resume_json)
    jobs_dir = st.text_input("岗位目录", default_jobs_dir)
    use_temp_weights = st.checkbox("使用临时权重在本地计算（不调用后端）", value=False, key="batch_use_local")
    run = st.button("开始批量匹配")
    if run:
        if not os.path.isfile(resume_json):
            st.error("简历JSON不存在")
            return
        if not os.path.isdir(jobs_dir):
            st.error("岗位目录不存在")
            return
        with open(resume_json, "r", encoding="utf-8") as f:
            resumes = json.load(f)
        jobs = []
        for name in os.listdir(jobs_dir):
            if not name.lower().endswith((".txt", ".md")):
                continue
            p = os.path.join(jobs_dir, name)
            try:
                with open(p, "r", encoding="utf-8") as jf:
                    jt = jf.read()
                jobs.append({"file": name, "text": jt})
            except Exception:
                pass
        results = []
        progress = st.progress(0)
        total = max(len(resumes) * max(len(jobs), 1), 1)
        done = 0
        for r in resumes:
            rtext = r.get("text", "")
            for j in jobs:
                if use_temp_weights and _matcher_mod is not None and "temp_weights" in st.session_state:
                    res = _local_match_with_weights(rtext, j.get("text", ""), st.session_state["temp_weights"])
                    results.append({
                        "resume_id": r.get("id"),
                        "job_file": j.get("file"),
                        "score": res.get("score"),
                        "details": res.get("details"),
                    })
                else:
                    payload = {"resume_text": rtext, "job_text": j.get("text", "")}
                    res, err = api_post(base, "/match", payload)
                    if err:
                        results.append({"resume_id": r.get("id"), "job_file": j.get("file"), "error": err})
                    else:
                        results.append({
                            "resume_id": r.get("id"),
                            "job_file": j.get("file"),
                            "score": res.get("score"),
                            "details": res.get("details"),
                        })
                done += 1
                progress.progress(min(done / total, 1.0))
        st.success(f"批量完成，共 {len(results)} 条结果")
        st.dataframe(results, use_container_width=True)
        st.download_button("下载JSON结果", data=json.dumps(results, ensure_ascii=False, indent=2), file_name="match_results.json")


def page_job_search():
    st.subheader("按岗位检索简历（Top N）")
    base = st.session_state.api_base
    # 岗位输入与上传
    c1, c2 = st.columns(2)
    with c1:
        job_text = st.text_area("岗位文本", height=160, placeholder="在此粘贴岗位描述…")
    with c2:
        jfile = st.file_uploader("上传岗位文件（txt/md/docx/pdf）", type=["txt", "md", "docx", "pdf"], key="job_search_file")
        if jfile is not None:
            job_text = _read_uploaded_text(jfile)
            st.info("已从上传文件填充岗位文本")

    source = st.radio("简历来源", ["使用处理后的JSON", "扫描简历目录"], index=0)
    top_k = st.number_input("Top N 输出", min_value=1, value=5, step=1)
    use_temp_weights = st.checkbox("使用临时权重本地计算（不调用后端）", value=False, key="job_search_local")

    resumes = []
    if source == "使用处理后的JSON":
        default_resume_json = os.path.join("data", "processed", "resumes_for_annotation.json")
        resume_json = st.text_input("简历JSON路径", default_resume_json, key="job_search_json")
        if st.button("计算并检索Top N", key="job_search_run_json"):
            if not job_text.strip():
                st.warning("请先填写岗位文本")
                return
            if not os.path.isfile(resume_json):
                st.error("简历JSON不存在")
                return
            with open(resume_json, "r", encoding="utf-8") as f:
                data = json.load(f)
            resumes = [{"id": (i if d.get("id") is None else d.get("id")), "text": d.get("text", "")} for i, d in enumerate(data)]
            _run_job_search(resumes, job_text, base, top_k, use_temp_weights)
    else:
        default_resume_dir = os.path.join("data", "raw_resumes")
        resume_dir = st.text_input("简历目录（docx/pdf/txt/md）", default_resume_dir, key="job_search_dir")
        if st.button("扫描目录并检索Top N", key="job_search_run_dir"):
            if not job_text.strip():
                st.warning("请先填写岗位文本")
                return
            if not os.path.isdir(resume_dir):
                st.error("简历目录不存在")
                return
            # 扫描目录并解析
            items = []
            for name in os.listdir(resume_dir):
                p = os.path.join(resume_dir, name)
                if not os.path.isfile(p):
                    continue
                low = name.lower()
                text = ""
                try:
                    if low.endswith((".txt", ".md")):
                        with open(p, "r", encoding="utf-8", errors="ignore") as f:
                            text = f.read()
                    elif low.endswith(".docx"):
                        if _preprocess_mod is not None:
                            text = _preprocess_mod.parse_word(p)
                    elif low.endswith(".pdf"):
                        if _preprocess_mod is not None:
                            text = _preprocess_mod.parse_pdf(p)
                except Exception:
                    text = ""
                if text:
                    items.append({"id": name, "text": text, "path": p})
            if not items:
                st.warning("目录中未解析到简历文本")
                return
            _run_job_search(items, job_text, base, top_k, use_temp_weights)

    # 如果已有检索结果（来自上一次运行），也在页面尾部展示，保证点击“查看”不会丢失数据
    if st.session_state.get("job_search_top"):
        st.markdown("---")
        _render_job_search_results(st.session_state["job_search_top"])

def page_ingest():
    st.subheader("在线采集与入库")
    base = st.session_state.api_base
    plat = st.selectbox("平台", ["通用URL", "Boss直聘", "智联招聘", "猎聘", "LinkedIn"], index=0, key="ingest_platform")
    urls_text = st.text_area("URL（每行一个）", height=120, key="ingest_urls")
    cookie = st.text_input("Cookie（需要账号登录后的Cookie）", value="", key="ingest_cookie")
    if st.button("采集并入库向量", key="ingest_run"):
        urls = [u.strip() for u in (urls_text or "").split("\n") if u.strip()]
        if not urls:
            st.warning("请填写至少一个URL")
            return
        if plat == "通用URL":
            res, err = _post_with_progress(base, "/ingest_online", {"urls": urls}, 12.0, "在线采集")
        elif plat == "Boss直聘":
            if not cookie.strip():
                st.error("Boss直聘需要Cookie")
                return
            res, err = _post_with_progress(base, "/ingest_bosszhipin", {"urls": urls, "cookie": cookie}, 12.0, "在线采集")
        elif plat == "智联招聘":
            if not cookie.strip():
                st.error("智联招聘需要Cookie")
                return
            res, err = _post_with_progress(base, "/ingest_zhilian", {"urls": urls, "cookie": cookie}, 12.0, "在线采集")
        elif plat == "猎聘":
            if not cookie.strip():
                st.error("猎聘需要Cookie")
                return
            res, err = _post_with_progress(base, "/ingest_liepin", {"urls": urls, "cookie": cookie}, 12.0, "在线采集")
        else:
            if not cookie.strip():
                st.error("LinkedIn需要Cookie")
                return
            res, err = _post_with_progress(base, "/ingest_linkedin", {"urls": urls, "cookie": cookie}, 12.0, "在线采集")
        if err:
            st.error(f"采集失败：{err}")
            return
        items = res.get("items", [])
        st.success(f"采集完成，得到 {len(items)} 条")
        rows = [{"url": it.get("url"), "text_len": len(it.get("text", "")), "tags": ",".join(it.get("fairness_tags", []))} for it in items]
        if rows:
            st.dataframe(rows, use_container_width=True)
        push = [{"id": str(i), "text": it.get("text", "")} for i, it in enumerate(items)]
        if push:
            _post_with_progress(base, "/vector_index", {"items": push}, 8.0, "向量入库")
            st.info("已入库到向量索引，可在‘三级漏斗’页面进行筛选")


def _run_job_search(resumes: List[Dict], job_text: str, base: str, top_k: int, use_temp_weights: bool):
    results = []
    progress = st.progress(0)
    total = max(len(resumes), 1)
    done = 0
    for r in resumes:
        rtext = r.get("text", "")
        if not rtext:
            continue
        if use_temp_weights and _matcher_mod is not None and "temp_weights" in st.session_state:
            res = _local_match_with_weights(rtext, job_text, st.session_state["temp_weights"])
            score = res.get("score", 0)
            details = res.get("details", {})
        else:
            res, err = api_post(base, "/match", {"resume_text": rtext, "job_text": job_text})
            if err:
                results.append({"resume_id": r.get("id"), "error": err})
                done += 1
                progress.progress(min(done / total, 1.0))
                continue
            score = res.get("score", 0)
            details = res.get("details", {})
        results.append({
            "resume_id": r.get("id"),
            "score": score,
            "details": details,
            # 便于点击查看原文内容与下载
            "resume_text": rtext,
            "resume_path": r.get("path"),
            "resume_profile": res.get("resume_profile", {}),
            "job_profile": res.get("job_profile", {})
        })
        done += 1
        progress.progress(min(done / total, 1.0))
    # 排序与Top N，并保存到会话状态，避免点击“查看”后丢失数据
    results = sorted([x for x in results if "score" in x], key=lambda x: x.get("score", 0), reverse=True)
    top = results[:top_k]
    st.session_state["job_search_top"] = top
    st.success(f"完成检索，共 {len(results)} 条结果；Top {top_k} 如下")


def _render_job_search_results(top: List[Dict]):
    if not top:
        st.info("暂无结果可展示")
        return
    # 仅展示关键列，避免在表格中显示大文本
    display_top = [{
        "resume_id": t.get("resume_id"),
        "score": t.get("score"),
        "matched_skills": ", ".join(t.get("details", {}).get("matched_skills", []))
    } for t in top]
    st.dataframe(display_top, use_container_width=True)
    st.download_button("下载TopN JSON", data=json.dumps(top, ensure_ascii=False, indent=2), file_name="top_resumes.json")

    # 交互：左右布局（左侧Top N列表，右侧对应简历原文）
    st.markdown("---")
    st.subheader("查看Top N的简历原文")
    ids = [t.get("resume_id") for t in top]
    default_id = ids[0] if ids else None
    if "top_view_select" not in st.session_state or st.session_state["top_view_select"] not in ids:
        st.session_state["top_view_select"] = default_id

    colL, colR = st.columns([1, 2])
    # 左侧：Top N 排序列表与选择
    with colL:
        rank_rows = [{"rank": i + 1, "resume_id": t.get("resume_id"), "score": t.get("score")} for i, t in enumerate(top)]
        st.dataframe(rank_rows, use_container_width=True)
        # 使用大按钮切换选择
        for i, t in enumerate(top):
            rid = t.get("resume_id")
            score = t.get("score")
            label = f"第{i+1}名：{rid}（分数 {score}）"
            if st.button(label, key=f"top_btn_{rid}", type="primary"):
                st.session_state["top_view_select"] = rid

    # 右侧：显示选中简历原文与下载 + 自动生成图表
    with colR:
        selected = st.session_state.get("top_view_select")
        if selected is not None:
            chosen = next((x for x in top if x.get("resume_id") == selected), None)
            if chosen:
                st.info(f"正在查看：{selected}")
                text = chosen.get("resume_text") or ""
                path = chosen.get("resume_path")
                if path and os.path.isfile(path):
                    low = path.lower()
                    try:
                        if low.endswith((".txt", ".md")):
                            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                                text = f.read()
                        elif low.endswith(".docx"):
                            parsed = ""
                            if _preprocess_mod is not None:
                                try:
                                    parsed = _preprocess_mod.parse_word(path)
                                except Exception:
                                    parsed = ""
                            if not parsed:
                                try:
                                    from docx import Document
                                    doc = Document(path)
                                    parsed = "\n".join(p.text for p in doc.paragraphs)
                                except Exception:
                                    parsed = ""
                            text = parsed or text
                        elif low.endswith(".pdf"):
                            parsed = ""
                            if _preprocess_mod is not None:
                                try:
                                    parsed = _preprocess_mod.parse_pdf(path)
                                except Exception:
                                    parsed = ""
                            if not parsed:
                                try:
                                    import PyPDF2
                                    with open(path, "rb") as pf:
                                        reader = PyPDF2.PdfReader(pf)
                                        for page in reader.pages:
                                            page_text = page.extract_text()
                                            if page_text:
                                                parsed += page_text
                                except Exception:
                                    parsed = ""
                            text = parsed or text
                    except Exception:
                        pass
                    try:
                        with open(path, "rb") as bf:
                            st.download_button("下载原文件", data=bf.read(), file_name=os.path.basename(path))
                    except Exception:
                        st.warning("无法提供原文件下载")

                st.text_area("简历内容预览", value=text, height=300)
                st.markdown("---")
                _render_match_charts(chosen.get("details", {}))
                _render_rich_features(chosen.get("details", {}), chosen.get("resume_profile", {}), chosen.get("job_profile", {}))

def _render_rich_features(details: Dict, rprof: Dict, jprof: Dict):
    try:
        import matplotlib.pyplot as plt
        th = _theme()
        fig, ax = plt.subplots(figsize=(5, 3))
        feats = [
            ("skills", float(details.get("skill_ratio", 0))),
            ("degree", float(details.get("degree_score", 0))),
            ("years", float(details.get("years_score", 0))),
            ("position", float(details.get("position_score", 0))),
            ("keywords", float(details.get("keyword_ratio", 0))),
            ("certs", float(details.get("certs_score", 0))),
            ("languages", float(details.get("languages_ratio", 0))),
        ]
        ax.bar([k for k, v in feats], [v for k, v in feats], color=th["primary"])
        ax.set_ylim(0, 1)
        ax.set_ylabel("score")
        ax.set_title("特征贡献")
        st.pyplot(fig, use_container_width=False)
    except Exception:
        pass
    # 饼图：技能匹配情况（命中 vs 未命中）
    try:
        import matplotlib.pyplot as plt
        th = _theme()
        job_sk = jprof.get("skills", [])
        matched = details.get("matched_skills", [])
        miss = max(len(job_sk) - len(matched), 0)
        fig2, ax2 = plt.subplots(figsize=(4, 3))
        ax2.pie([len(matched), miss], labels=["命中", "未命中"], autopct='%1.0f%%', colors=[th["success"], th["warning"]])
        ax2.set_title("技能命中率")
        st.pyplot(fig2, use_container_width=False)
    except Exception:
        pass

def page_config_view():
    st.subheader("匹配配置查看")
    cfg_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "config", "matching.json")
    st.write("配置文件路径：", cfg_path)
    if os.path.isfile(cfg_path):
        with open(cfg_path, "r", encoding="utf-8") as f:
            cfg = json.load(f)
        st.json(cfg)
        st.info("提示：当前API在导入匹配脚本时加载配置；如需修改配置生效，请重启后端。")
    else:
        st.warning("未找到配置文件。")

    st.markdown("---")
    st.subheader("临时权重试验（不改文件，仅前端生效）")
    c1, c2, c3, c4 = st.columns(4)
    w_skills = c1.slider("skills", 0.0, 1.0, 0.5, 0.05)
    w_degree = c2.slider("degree", 0.0, 1.0, 0.2, 0.05)
    w_years = c3.slider("years", 0.0, 1.0, 0.2, 0.05)
    w_position = c4.slider("position", 0.0, 1.0, 0.1, 0.05)
    if st.button("应用到本地匹配"):
        st.session_state["temp_weights"] = {
            "skills": float(w_skills),
            "degree": float(w_degree),
            "years": float(w_years),
            "position": float(w_position),
        }
        st.success("已设置临时权重（仅前端本地匹配有效）")
    st.markdown("---")
    st.subheader("行业模板库编辑")
    base = st.session_state.api_base
    tmpl = {}
    res, err = api_get(base, "/config/industry_templates", timeout=2.5)
    if not err and res:
        tmpl = res.get("templates", {})
    else:
        tmpl = {}
    edit_text = st.text_area("编辑JSON", value=json.dumps(tmpl, ensure_ascii=False, indent=2), height=200)
    if st.button("保存行业模板"):
        try:
            new_obj = json.loads(edit_text)
        except Exception:
            st.error("JSON不合法")
            new_obj = None
        if new_obj is not None:
            res, err = api_post(base, "/config/industry_templates", {"templates": new_obj})
            if err:
                st.error(f"保存失败：{err}")
            else:
                st.success("已保存并应用行业模板")


def _read_uploaded_text(file) -> str:
    name = (file.name or "").lower()
    try:
        if name.endswith((".txt", ".md")):
            return file.read().decode("utf-8", errors="ignore")
        if name.endswith(".docx"):
            try:
                from docx import Document
                doc = Document(file)
                return "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                # 回退到预处理模块方法（若存在，需要写入临时文件）
                if _preprocess_mod is not None:
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                        tmp.write(file.read())
                        tmp_path = tmp.name
                    try:
                        return _preprocess_mod.parse_word(tmp_path)
                    finally:
                        try:
                            os.remove(tmp_path)
                        except Exception:
                            pass
                return ""
        if name.endswith(".pdf"):
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
                return text
            except Exception:
                # 回退到预处理模块方法（若存在，需要写入临时文件）
                if _preprocess_mod is not None:
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                        tmp.write(file.read())
                        tmp_path = tmp.name
                    try:
                        return _preprocess_mod.parse_pdf(tmp_path)
                    finally:
                        try:
                            os.remove(tmp_path)
                        except Exception:
                            pass
                return ""
        # 其他类型按文本尝试
        return file.read().decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _local_match_with_weights(resume_text: str, job_text: str, weights: Dict[str, float]) -> Dict:
    if _matcher_mod is None:
        return {"score": 0, "details": {"error": "matcher module not available"}}
    # 覆盖权重（仅本次计算，不持久化）
    orig = dict(_matcher_mod.WEIGHTS)
    new_w = dict(orig)
    new_w.update({k: float(v) for k, v in weights.items()})
    _matcher_mod.WEIGHTS = new_w
    try:
        rprofile = _matcher_mod.profile_from_text(resume_text)
        jprofile = _matcher_mod.job_profile_from_text(job_text)
        score, details = _matcher_mod.match_score(rprofile, jprofile)
        return {
            "score": score,
            "resume_profile": rprofile,
            "job_profile": jprofile,
            "details": details,
        }
    finally:
        _matcher_mod.WEIGHTS = orig


def _render_match_result(res: Dict):
    st.success("匹配完成")
    score = res.get("score", 0)
    details = res.get("details", {})
    colA, colB, colC = st.columns(3)
    colA.metric("综合得分", f"{score}")
    colB.metric("技能匹配比例", f"{details.get('skill_ratio', 0)}")
    colC.metric("年限匹配", f"{details.get('years_score', 0)}")
    st.write("命中技能：", ", ".join(details.get("matched_skills", [])))
    st.write("学历匹配：", details.get("degree_score", 0))
    st.write("职位匹配：", details.get("position_score", 0))
    st.expander("简历画像").json(res.get("resume_profile", {}))
    st.expander("岗位画像").json(res.get("job_profile", {}))
    _render_match_charts(details)
    _render_match_pipeline(details, float(score))


def _render_match_charts(details: Dict):
    dims = [
        ("技能", float(details.get("skill_ratio", 0))),
        ("学历", float(details.get("degree_score", 0))),
        ("年限", float(details.get("years_score", 0))),
        ("职位", float(details.get("position_score", 0))),
        ("关键词", float(details.get("keyword_ratio", 0))),
        ("证书", float(details.get("certs_score", 0))),
        ("语言", float(details.get("languages_ratio", 0))),
        ("格式", float(details.get("format_score", 0))),
    ]
    if st_echarts:
        ind = [{"name": k, "max": 1.0} for k, _ in dims]
        val = [v for _, v in dims]
        opts = {
            "legend": {"data": ["匹配维度"]},
            "radar": {"indicator": ind},
            "series": [{"type": "radar", "data": [{"value": val, "name": "匹配维度"}]}]
        }
        st_echarts(opts, height=360)
        matched = len(details.get("matched_skills", []))
        total_sk = int(details.get("total_job_skills", matched) or matched)
        miss = max(total_sk - matched, 0)
        p_opts = {
            "tooltip": {"trigger": "item"},
            "series": [{
                "type": "pie",
                "radius": "50%",
                "data": [
                    {"value": matched, "name": "命中"},
                    {"value": miss, "name": "未命中"}
                ]
            }]
        }
        st_echarts(p_opts, height=300)
    else:
        import pandas as pd
        st.bar_chart(pd.DataFrame({"score": [v for _, v in dims]}, index=[k for k, _ in dims]))
        matched = len(details.get("matched_skills", []))
        miss = max(int(details.get("total_job_skills", matched) or matched) - matched, 0)
        st.write("技能命中/未命中：", matched, "/", miss)

def _render_candidate_insights(details: Dict, rprof: Dict, jprof: Dict, resume_text: str, job_text: str, base: str):
    _render_match_charts(details)
    st.markdown("**能力结构占比**")
    contrib = [
        ("硬性", (float(details.get("degree_score", 0)) + float(details.get("years_score", 0)) + float(details.get("position_score", 0))) / 3.0),
        ("软性", float(details.get("skill_ratio", 0))),
        ("格式", float(details.get("format_score", 0)))
    ]
    if st_echarts:
        p = {
            "tooltip": {"trigger": "item"},
            "series": [{"type": "pie", "radius": "55%", "data": [{"name": k, "value": v} for k, v in contrib]}]
        }
        st_echarts(p, height=300)
    else:
        import pandas as pd
        st.bar_chart(pd.DataFrame({"value": [v for _, v in contrib]}, index=[k for k, _ in contrib]))
    # 面试题与分析由外层页面统一排版与生成，避免此处重复
    try:
        import pandas as pd, io
        cols = ["skill_ratio","degree_score","years_score","position_score","keyword_ratio","certs_score","languages_ratio","format_score"]
        df1 = pd.DataFrame([{"维度": c, "分数": float(details.get(c, 0))} for c in cols])
        ms = details.get("matched_skills", []) or []
        df2 = pd.DataFrame([{"技能": s} for s in ms])
        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine="xlsxwriter") as w:
            df1.to_excel(w, index=False, sheet_name="维度分数")
            df2.to_excel(w, index=False, sheet_name="命中技能")
        st.download_button("导出候选人详情Excel", data=buf.getvalue(), file_name="candidate_detail.xlsx")
    except Exception:
        pass

def _render_match_pipeline(details: Dict, final_score: float):
    hard = (
        float(details.get("degree_score", 0)) +
        float(details.get("years_score", 0)) +
        float(details.get("position_score", 0))
    ) / 3.0
    soft = float(details.get("skill_ratio", 0))
    fmt = float(details.get("format_score", 0))
    comp = 0.4 * hard + 0.5 * soft + 0.1 * fmt
    st.markdown("---")
    st.subheader("阶段评分")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("硬性条件", f"{round(hard, 4)}")
    c2.metric("软性条件", f"{round(soft, 4)}")
    c3.metric("格式规范", f"{round(fmt, 4)}")
    c4.metric("综合加权", f"{round(comp, 4)}")
    st.subheader("综合评价与建议")
    adv = []
    if hard < 0.5:
        adv.append("硬性条件不足，可优先筛选学历/年限更匹配候选")
    if soft < 0.5:
        adv.append("技能匹配较弱，建议关注核心技能训练或转岗可能")
    if fmt < 0.4:
        adv.append("简历格式较弱，建议优化结构与关键经历描述")
    if final_score >= 0.8:
        adv.append("总体高度匹配，建议进入面试流程")
    elif final_score >= 0.6:
        adv.append("总体较为匹配，可考虑电话初筛后决定")
    else:
        adv.append("总体匹配一般，建议放入储备或不进入下一轮")
    st.write("；".join(adv))


def page_ner_eval():
    st.subheader("NER 评估（实体级 Precision/Recall/F1）")
    base = st.session_state.api_base
    data_path = st.text_input("标注数据路径", os.path.join("data", "processed", "entity_train.json"))
    sample_n = st.number_input("评估样本数（0 为全部）", min_value=0, value=0, step=10)
    run = st.button("开始评估")
    if run:
        if not os.path.isfile(data_path):
            st.error("标注数据不存在")
            return
        with open(data_path, "r", encoding="utf-8") as f:
            dataset = json.load(f)
        if sample_n and sample_n > 0:
            dataset = dataset[:sample_n]
        TP = FP = FN = 0
        rows = []
        trace_prec = []
        trace_rec = []
        trace_f1 = []
        details_rows = []
        import collections
        type_gt = collections.Counter()
        type_pred = collections.Counter()
        cum_type = collections.defaultdict(lambda: {"TP": 0, "FP": 0, "FN": 0})
        trace_type = collections.defaultdict(lambda: {"prec": [], "rec": [], "f1": []})
        fp_examples = []
        fn_examples = []
        prog = st.progress(0.0)
        total = max(len(dataset), 1)
        for i, item in enumerate(dataset):
            text = item.get("text", "")
            gt = {(e.get("type"), e.get("text")) for e in item.get("entities", []) if e.get("type") and e.get("text")}
            res, err = api_post(base, "/predict", {"text": text})
            if err:
                rows.append({"idx": i, "error": err})
                continue
            pred = {(e.get("type"), e.get("text")) for e in res.get("entities", []) if e.get("type") and e.get("text")}
            tp = len(gt & pred)
            fp = len(pred - gt)
            fn = len(gt - pred)
            TP += tp
            FP += fp
            FN += fn
            prec_i = tp / (tp + fp) if (tp + fp) > 0 else 0.0
            rec_i = tp / (tp + fn) if (tp + fn) > 0 else 0.0
            f1_i = (2 * prec_i * rec_i) / (prec_i + rec_i) if (prec_i + rec_i) > 0 else 0.0
            rows.append({"idx": i, "tp": tp, "fp": fp, "fn": fn, "gt": len(gt), "pred": len(pred), "precision": round(prec_i, 4), "recall": round(rec_i, 4), "f1": round(f1_i, 4)})
            trace_prec.append((i + 1, (TP / (TP + FP)) if (TP + FP) > 0 else 0.0))
            trace_rec.append((i + 1, (TP / (TP + FN)) if (TP + FN) > 0 else 0.0))
            trace_f1.append((i + 1, ((2 * (TP / (TP + FP)) * (TP / (TP + FN))) / ((TP / (TP + FP)) + (TP / (TP + FN)))) if ((TP + FP) > 0 and (TP + FN) > 0 and ((TP / (TP + FP)) + (TP / (TP + FN))) > 0) else 0.0))
            details_rows.append({"idx": i, "gt": sorted(list(gt)), "pred": sorted(list(pred))})
            for e in item.get("entities", []) or []:
                t = e.get("type")
                if t:
                    type_gt[t] += 1
            for e in res.get("entities", []) or []:
                t = e.get("type")
                if t:
                    type_pred[t] += 1
            for e in (pred - gt):
                fp_examples.append({"idx": i, "type": e[0], "text": e[1]})
            for e in (gt - pred):
                fn_examples.append({"idx": i, "type": e[0], "text": e[1]})
            tset = sorted(set([x[0] for x in gt] + [x[0] for x in pred]))
            for t in tset:
                tp_t = len([x for x in (gt & pred) if x[0] == t])
                fp_t = len([x for x in (pred - gt) if x[0] == t])
                fn_t = len([x for x in (gt - pred) if x[0] == t])
                cum_type[t]["TP"] += tp_t
                cum_type[t]["FP"] += fp_t
                cum_type[t]["FN"] += fn_t
                ctp = cum_type[t]["TP"]
                cfp = cum_type[t]["FP"]
                cfn = cum_type[t]["FN"]
                p_t = ctp / (ctp + cfp) if (ctp + cfp) > 0 else 0.0
                r_t = ctp / (ctp + cfn) if (ctp + cfn) > 0 else 0.0
                f_t = (2 * p_t * r_t) / (p_t + r_t) if (p_t + r_t) > 0 else 0.0
                trace_type[t]["prec"].append((i + 1, p_t))
                trace_type[t]["rec"].append((i + 1, r_t))
                trace_type[t]["f1"].append((i + 1, f_t))
            prog.progress(min((i + 1) / float(total), 1.0))
        precision = TP / (TP + FP) if (TP + FP) > 0 else 0.0
        recall = TP / (TP + FN) if (TP + FN) > 0 else 0.0
        f1 = (2 * precision * recall) / (precision + recall) if (precision + recall) > 0 else 0.0
        st.metric("Precision", f"{precision:.4f}")
        st.metric("Recall", f"{recall:.4f}")
        st.metric("F1", f"{f1:.4f}")
        st.dataframe(rows, use_container_width=True)
        st.markdown("---")
        st.subheader("过程曲线")
        try:
            th = _theme()
            if st_echarts:
                xs = [x for x, y in trace_prec]
                opts = {
                    "tooltip": {"trigger": "axis"},
                    "xAxis": {"type": "category", "data": xs, "name": "样本序号"},
                    "yAxis": {"type": "value", "name": "metric"},
                    "series": [
                        {"type": "line", "name": "Precision", "data": [y for x, y in trace_prec], "smooth": True, "itemStyle": {"color": th["primary"]}},
                        {"type": "line", "name": "Recall", "data": [y for x, y in trace_rec], "smooth": True, "itemStyle": {"color": th["success"]}},
                        {"type": "line", "name": "F1", "data": [y for x, y in trace_f1], "smooth": True, "itemStyle": {"color": th["warning"]}},
                    ]
                }
                st_echarts(opts, height=320)
            else:
                import matplotlib.pyplot as plt
                fig, ax = plt.subplots(figsize=(6, 3))
                xs = [x for x, y in trace_prec]
                ax.plot(xs, [y for x, y in trace_prec], label="Precision")
                ax.plot(xs, [y for x, y in trace_rec], label="Recall")
                ax.plot(xs, [y for x, y in trace_f1], label="F1")
                ax.set_ylim(0, 1)
                ax.legend()
                st.pyplot(fig, use_container_width=False)
        except Exception:
            pass
        try:
            import matplotlib.pyplot as plt
            fig, ax = plt.subplots(figsize=(5, 3))
            th = _theme()
            ax.bar(["TP", "FP", "FN"], [TP, FP, FN], color=[th["success"], th["danger"], th["warning"]])
            st.pyplot(fig, use_container_width=False)
        except Exception:
            pass
        st.markdown("---")
        st.subheader("类型分布")
        try:
            import matplotlib.pyplot as plt
            cats = sorted(set(list(type_gt.keys()) + list(type_pred.keys())))
            vals_gt = [type_gt.get(c, 0) for c in cats]
            vals_pd = [type_pred.get(c, 0) for c in cats]
            th = _theme()
            fig3, ax3 = plt.subplots(figsize=(6, 3))
            w = 0.35
            xs = list(range(len(cats)))
            ax3.bar([x - w/2 for x in xs], vals_gt, width=w, color=th["primary"], label="GT")
            ax3.bar([x + w/2 for x in xs], vals_pd, width=w, color=th["success"], label="Pred")
            ax3.set_xticks(xs)
            ax3.set_xticklabels(cats, rotation=30)
            ax3.legend()
            st.pyplot(fig3, use_container_width=False)
        except Exception:
            pass
        st.markdown("---")
        st.subheader("按类型指标与曲线")
        cats = sorted(set(list(type_gt.keys()) + list(type_pred.keys()) + list(cum_type.keys())))
        type_rows = []
        for t in cats:
            ctp = cum_type[t]["TP"]
            cfp = cum_type[t]["FP"]
            cfn = cum_type[t]["FN"]
            p_t = ctp / (ctp + cfp) if (ctp + cfp) > 0 else 0.0
            r_t = ctp / (ctp + cfn) if (ctp + cfn) > 0 else 0.0
            f_t = (2 * p_t * r_t) / (p_t + r_t) if (p_t + r_t) > 0 else 0.0
            type_rows.append({"type": t, "TP": ctp, "FP": cfp, "FN": cfn, "precision": round(p_t, 4), "recall": round(r_t, 4), "f1": round(f_t, 4)})
        if type_rows:
            import pandas as pd
            df_types = pd.DataFrame(type_rows)
            st.dataframe(df_types, use_container_width=True)
            st.download_button("下载按类型指标CSV", data=df_types.to_csv(index=False).encode("utf-8-sig"), file_name="ner_type_metrics.csv")
            sel_type = st.selectbox("选择实体类型查看曲线", cats, index=0)
            tr = trace_type.get(sel_type)
            if tr and (tr.get("prec") or tr.get("rec") or tr.get("f1")):
                try:
                    th = _theme()
                    if st_echarts:
                        xs = [x for x, y in tr.get("prec", [])]
                        opts_t = {
                            "tooltip": {"trigger": "axis"},
                            "xAxis": {"type": "category", "data": xs, "name": "样本序号"},
                            "yAxis": {"type": "value", "name": "metric"},
                            "series": [
                                {"type": "line", "name": "Precision", "data": [y for x, y in tr.get("prec", [])], "smooth": True, "itemStyle": {"color": th["primary"]}},
                                {"type": "line", "name": "Recall", "data": [y for x, y in tr.get("rec", [])], "smooth": True, "itemStyle": {"color": th["success"]}},
                                {"type": "line", "name": "F1", "data": [y for x, y in tr.get("f1", [])], "smooth": True, "itemStyle": {"color": th["warning"]}},
                            ]
                        }
                        st_echarts(opts_t, height=320)
                    else:
                        import matplotlib.pyplot as plt
                        figt, axt = plt.subplots(figsize=(6, 3))
                        xs = [x for x, y in tr.get("prec", [])]
                        axt.plot(xs, [y for x, y in tr.get("prec", [])], label="Precision")
                        axt.plot(xs, [y for x, y in tr.get("rec", [])], label="Recall")
                        axt.plot(xs, [y for x, y in tr.get("f1", [])], label="F1")
                        axt.set_ylim(0, 1)
                        axt.legend()
                        st.pyplot(figt, use_container_width=False)
                except Exception:
                    pass
        st.markdown("---")
        st.subheader("错例浏览与导出")
        if fp_examples:
            import pandas as pd
            df_fp = pd.DataFrame(fp_examples)
            st.write("FP（预测多余）")
            st.dataframe(df_fp, use_container_width=True)
            st.download_button("下载FP JSON", data=json.dumps(fp_examples, ensure_ascii=False, indent=2), file_name="ner_fp.json")
            st.download_button("下载FP CSV", data=df_fp.to_csv(index=False).encode("utf-8-sig"), file_name="ner_fp.csv")
        if fn_examples:
            import pandas as pd
            df_fn = pd.DataFrame(fn_examples)
            st.write("FN（预测缺失）")
            st.dataframe(df_fn, use_container_width=True)
            st.download_button("下载FN JSON", data=json.dumps(fn_examples, ensure_ascii=False, indent=2), file_name="ner_fn.json")
            st.download_button("下载FN CSV", data=df_fn.to_csv(index=False).encode("utf-8-sig"), file_name="ner_fn.csv")
        st.markdown("---")
        st.subheader("评估报告导出")
        import io
        report_html = f"""
        <html><head><meta charset='utf-8'><title>NER评估报告</title></head><body>
        <h2>汇总指标</h2>
        <ul>
        <li>Precision: {precision:.4f}</li>
        <li>Recall: {recall:.4f}</li>
        <li>F1: {f1:.4f}</li>
        </ul>
        <h2>类型指标</h2>
        <table border='1' cellspacing='0' cellpadding='6'>
        <tr><th>type</th><th>TP</th><th>FP</th><th>FN</th><th>precision</th><th>recall</th><th>f1</th></tr>
        {''.join([f"<tr><td>{r['type']}</td><td>{r['TP']}</td><td>{r['FP']}</td><td>{r['FN']}</td><td>{r['precision']}</td><td>{r['recall']}</td><td>{r['f1']}</td></tr>" for r in type_rows])}
        </table>
        </body></html>
        """
        st.download_button("下载评估报告HTML", data=report_html.encode("utf-8"), file_name="ner_eval_report.html")
        try:
            import matplotlib.pyplot as plt
            buf_proc = io.BytesIO()
            xs = [x for x, y in trace_prec]
            figp, axp = plt.subplots(figsize=(6, 3))
            axp.plot(xs, [y for x, y in trace_prec], label="Precision")
            axp.plot(xs, [y for x, y in trace_rec], label="Recall")
            axp.plot(xs, [y for x, y in trace_f1], label="F1")
            axp.set_ylim(0, 1)
            axp.legend()
            figp.savefig(buf_proc, format="png", dpi=160)
            st.download_button("下载过程曲线PNG", data=buf_proc.getvalue(), file_name="ner_proc_curves.png")
        except Exception:
            pass
        try:
            import matplotlib.pyplot as plt
            buf_conf = io.BytesIO()
            th = _theme()
            figc, axc = plt.subplots(figsize=(5, 3))
            axc.bar(["TP", "FP", "FN"], [TP, FP, FN], color=[th["success"], th["danger"], th["warning"]])
            figc.savefig(buf_conf, format="png", dpi=160)
            st.download_button("下载混淆计数PNG", data=buf_conf.getvalue(), file_name="ner_confusion_counts.png")
        except Exception:
            pass
        try:
            import matplotlib.pyplot as plt
            buf_types = io.BytesIO()
            cats = sorted(set(list(type_gt.keys()) + list(type_pred.keys())))
            vals_gt = [type_gt.get(c, 0) for c in cats]
            vals_pd = [type_pred.get(c, 0) for c in cats]
            th = _theme()
            figt3, axt3 = plt.subplots(figsize=(6, 3))
            w = 0.35
            xs = list(range(len(cats)))
            axt3.bar([x - w/2 for x in xs], vals_gt, width=w, color=th["primary"], label="GT")
            axt3.bar([x + w/2 for x in xs], vals_pd, width=w, color=th["success"], label="Pred")
            axt3.set_xticks(xs)
            axt3.set_xticklabels(cats, rotation=30)
            axt3.legend()
            figt3.savefig(buf_types, format="png", dpi=160)
            st.download_button("下载类型分布PNG", data=buf_types.getvalue(), file_name="ner_type_distribution.png")
        except Exception:
            pass
        st.markdown("---")
        st.subheader("样本详情")
        if details_rows:
            sel = st.number_input("样本序号", min_value=0, max_value=len(details_rows)-1, value=0, step=1)
            chosen = details_rows[int(sel)]
            st.write("GT 实体：")
            st.dataframe([{"type": t, "text": v} for t, v in chosen.get("gt", [])], use_container_width=True)
            st.write("Pred 实体：")
            st.dataframe([{"type": t, "text": v} for t, v in chosen.get("pred", [])], use_container_width=True)


def _api_upload(base: str, file_obj, text: str | None, filename: str | None):
    import requests
    url = base.rstrip("/") + "/uploads"
    files = None
    data = {}
    if file_obj is not None:
        name = filename or getattr(file_obj, "name", None) or "upload.bin"
        buf = file_obj.read()
        files = {"file": (name, buf)}
        data["filename"] = name
    elif text is not None and text.strip():
        data["text"] = text
        if filename:
            data["filename"] = filename
    r = requests.post(url, files=files, data=data, timeout=float(st.session_state.get("api_timeout", 12.0)))
    return r.json()

def page_candidate_upload():
    st.subheader("简历上传")
    base = st.session_state.api_base
    f = st.file_uploader("选择简历文件", type=["txt", "md", "docx", "pdf"]) 
    text = st.text_area("或粘贴简历文本", height=180)
    name = st.text_input("文件名(可选)")
    run = st.button("提交")
    if run:
        try:
            res = _api_upload(base, f, text, name)
            st.success(f"上传成功：{res.get('path')} 大小 {res.get('size')}")
        except Exception as e:
            st.error(str(e))

def _role_pages() -> list[tuple[str, callable]]:
    r = st.session_state.get("role")
    all_pages = [
        ("实体预测", page_predict),
        ("单次匹配", page_match_single),
        ("批量匹配", page_match_batch),
        ("岗位检索", page_job_search),
        ("在线采集与入库", page_ingest),
        ("三级漏斗", page_funnel),
        ("公平性报告", page_fairness),
        ("决策辅助", page_decision),
        ("配置查看", page_config_view),
        ("简历上传", page_candidate_upload),
        ("NER评估", page_ner_eval),
    ]
    if r == "管理员":
        return [("管理员流程向导", page_admin_wizard)]
    if r == "面向招聘方":
        return [("招聘流程向导", page_hr_wizard)]
    if r == "面向求职者":
        return [("求职者流程向导", page_jobseeker_wizard)]
    return all_pages

def main():
    page_setup()
    sidebar()
    pages = _role_pages()
    tabs = st.tabs([t for t, _ in pages])
    for i, (_, fn) in enumerate(pages):
        with tabs[i]:
            fn()

def service_overview():
    st.subheader("服务概览")
    base = st.session_state.api_base
    ok = try_health(base.rstrip("/") + "/health")
    st.metric("API健康", "正常" if ok else "异常")
    st.write("API地址:", base)
    cfg_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "config", "matching.json")
    if os.path.isfile(cfg_path):
        with open(cfg_path, "r", encoding="utf-8") as f:
            try:
                import json
                cfg = json.load(f)
                st.expander("匹配配置预览").json(cfg)
            except Exception:
                st.info("配置读取失败")

def page_jd_generate():
    st.subheader("JD智能生成")
    base = st.session_state.api_base
    text = st.text_area("用自然语言描述招聘需求", height=200)
    run = st.button("生成结构化JD")
    if run and text.strip():
        res, err = api_post(base, "/jd_generate", {"text": text}, timeout=6.0)
        if err:
            st.error(err)
        else:
            obj = res.get("jd", {})
            st.json(obj)
            import pandas as pd, io
            rows = []
            for k, v in obj.items():
                if isinstance(v, list):
                    rows.append({"field": k, "value": "; ".join([str(x) for x in v])})
                else:
                    rows.append({"field": k, "value": str(v)})
            st.dataframe(rows, use_container_width=True)
            try:
                buf = io.BytesIO()
                with pd.ExcelWriter(buf, engine="xlsxwriter") as w:
                    pd.DataFrame(rows).to_excel(w, index=False, sheet_name="JD")
                st.download_button("下载JD结构Excel", data=buf.getvalue(), file_name="structured_jd.xlsx")
            except Exception:
                csv = pd.DataFrame(rows).to_csv(index=False).encode("utf-8-sig")
                st.download_button("下载JD结构CSV", data=csv, file_name="structured_jd.csv")

def page_hr_wizard():
    base = st.session_state.api_base
    st.markdown("<div class='step-title'>步骤1：岗位需求</div>", unsafe_allow_html=True)
    st.caption("先上传JD文件或输入文本，生成结构化JD；可临时调节匹配权重")
    jd_file = st.file_uploader("上传JD文件（txt/md/docx/pdf）", type=["txt", "md", "docx", "pdf"], key="hr_jd_file")
    if jd_file is not None:
        try:
            st.session_state["hr_jd_preview"] = (_read_uploaded_text(jd_file) or "")[:800]
        except Exception:
            st.session_state["hr_jd_preview"] = ""
    if st.button("填入示例JD", key="hr_jd_example", type="secondary"):
        example = "招聘AI开发工程师，负责模型训练与部署，要求本科及以上，3年以上相关经验，熟悉Python、Docker、K8s，具备云平台经验（AWS/Azure/GCP）。"
        st.session_state["hr_jd_text"] = example
        st.session_state["hr_jd_preview"] = example
    jd_text = st.text_area("自然语言JD描述", height=160, key="hr_jd_text")
    gen_jd = st.button("生成结构化JD", key="hr_gen_jd")
    if gen_jd:
        if jd_file is not None and not jd_text.strip():
            jd_text = jd_file.read().decode("utf-8", errors="ignore")
        res, err = api_post(base, "/jd_generate", {"text": jd_text or ""}, timeout=6.0)
        if err:
            st.error(err)
        else:
            st.session_state["hr_struct_jd"] = res.get("jd", {})
            st.success("已生成结构化JD")
    if st.session_state.get("hr_struct_jd"):
        obj = st.session_state["hr_struct_jd"]
        import pandas as pd
        cn = {
            "degree_required": "学历要求",
            "min_years": "最低年限",
            "salary_min": "最低薪资",
            "salary_max": "最高薪资",
            "employment_type": "雇佣类型",
            "location": "地点",
            "industry": "行业",
            "skills": "技能",
            "frameworks": "框架",
            "tools": "工具",
            "languages": "编程语言",
            "soft_skills": "软技能",
            "keywords": "关键词",
            "certifications": "证书"
        }
        hard_keys = ["degree_required","min_years","salary_min","salary_max","employment_type","location","industry"]
        soft_keys = ["skills","frameworks","tools","languages","soft_skills"]
        entity_keys = ["keywords","certifications"]
        hard = {cn.get(k, k): obj.get(k) for k in hard_keys}
        soft = {cn.get(k, k): obj.get(k) for k in soft_keys}
        entity = {cn.get(k, k): obj.get(k) for k in entity_keys}
        preview = st.session_state.get("hr_jd_preview", "")
        tabs = st.tabs(["表格", "JSON"])
        with tabs[0]:
            st.subheader("硬性要求")
            st.dataframe(pd.DataFrame([{"字段": k, "值": str(v) if not isinstance(v, list) else "; ".join(str(x) for x in v or [])} for k, v in hard.items()]), use_container_width=True)
            st.subheader("软性要求")
            st.dataframe(pd.DataFrame([{"字段": k, "值": str(v) if not isinstance(v, list) else "; ".join(str(x) for x in v or [])} for k, v in soft.items()]), use_container_width=True)
            st.subheader("实体特征")
            ent_rows = [{"字段": k, "值": str(v) if not isinstance(v, list) else "; ".join(str(x) for x in v or [])} for k, v in entity.items()]
            if preview:
                ent_rows.insert(0, {"字段": "预览", "值": preview[:300]})
            st.dataframe(pd.DataFrame(ent_rows), use_container_width=True)
            if preview:
                with st.expander("JD预览文本"):
                    st.text_area("预览", value=preview, height=180, key="hr_jd_preview_area")
        with tabs[1]:
            st.json(obj)
        wcol = st.columns(4)
        with wcol[0]:
            st.write("匹配权重临时调节")
        keys = ["skill_ratio","degree_score","years_score","position_score","keyword_ratio","certs_score","languages_ratio","format_score"]
        sliders = {}
        for k in keys:
            sliders[k] = st.slider(k, 0.0, 1.0, float(st.session_state.get("hr_weights", {}).get(k, 0.1 if k!="skill_ratio" else 0.3)), 0.05, key=f"hr_w_{k}")
        st.session_state["hr_weights"] = sliders
    st.markdown("<div class='step-title'>步骤2：简历来源/在线采集</div>", unsafe_allow_html=True)
    st.caption("选择目录解析或在线采集合并岗位池，形成候选集合")
    t1, t2 = st.tabs(["简历来源", "在线采集合并岗位池"])
    with t1:
        mode = st.radio("来源", ["目录", "单文件"], index=0, key="hr_src_mode")
        if mode == "目录":
            resume_dir = st.text_input("简历目录", os.path.join("data", "raw_resumes"), key="hr_resume_dir")
            scan = st.button("扫描并解析目录", key="hr_scan_dir")
            if scan:
                if not os.path.isdir(resume_dir):
                    st.error("目录不存在")
                else:
                    items = []
                    for name in os.listdir(resume_dir):
                        p = os.path.join(resume_dir, name)
                        if not os.path.isfile(p):
                            continue
                        low = name.lower()
                        text = ""
                        try:
                            if low.endswith((".txt", ".md")):
                                with open(p, "r", encoding="utf-8", errors="ignore") as f:
                                    text = f.read()
                            elif low.endswith(".docx") and _preprocess_mod is not None:
                                text = _preprocess_mod.parse_word(p)
                            elif low.endswith(".pdf") and _preprocess_mod is not None:
                                text = _preprocess_mod.parse_pdf(p)
                        except Exception:
                            text = ""
                        if text:
                            items.append({"id": name, "text": text})
                    st.session_state["hr_resumes"] = items
                    st.success(f"已解析 {len(items)} 份简历")
        else:
            rf = st.file_uploader("上传简历文件（txt/md/docx/pdf）", type=["txt", "md", "docx", "pdf"], key="hr_resume_file")
            if st.button("解析文件", key="hr_parse_file"):
                if rf is None:
                    st.warning("请上传文件")
                else:
                    text = _read_uploaded_text(rf)
                    if text:
                        st.session_state["hr_resumes"] = [{"id": getattr(rf, "name", "resume"), "text": text}]
                        st.success("已解析 1 份简历")
        resumes = st.session_state.get("hr_resumes") or []
    with t2:
        plat = st.selectbox("平台", ["通用URL", "Boss直聘", "智联招聘", "猎聘", "LinkedIn"], index=0, key="hr_ingest_platform")
        urls_text = st.text_area("URL（每行一个）", height=120, key="hr_ingest_urls")
        cookie = st.text_input("Cookie（需登录后Cookie）", value="", key="hr_ingest_cookie")
        do_ingest = st.button("采集并合并到岗位池", key="hr_ingest_run")
        if do_ingest:
            urls = [u.strip() for u in (urls_text or "").split("\n") if u.strip()]
            if urls:
                if plat == "通用URL":
                    res, err = _post_with_progress(base, "/ingest_online", {"urls": urls}, 12.0, "在线采集")
                elif plat == "Boss直聘":
                    if not cookie.strip():
                        st.error("Boss直聘需要Cookie")
                        res, err = (None, "cookie required")
                    else:
                        res, err = _post_with_progress(base, "/ingest_bosszhipin", {"urls": urls, "cookie": cookie}, 12.0, "在线采集")
                elif plat == "智联招聘":
                    if not cookie.strip():
                        st.error("智联招聘需要Cookie")
                        res, err = (None, "cookie required")
                    else:
                        res, err = _post_with_progress(base, "/ingest_zhilian", {"urls": urls, "cookie": cookie}, 12.0, "在线采集")
                elif plat == "猎聘":
                    if not cookie.strip():
                        st.error("猎聘需要Cookie")
                        res, err = (None, "cookie required")
                    else:
                        res, err = _post_with_progress(base, "/ingest_liepin", {"urls": urls, "cookie": cookie}, 12.0, "在线采集")
                else:
                    if not cookie.strip():
                        st.error("LinkedIn需要Cookie")
                        res, err = (None, "cookie required")
                    else:
                        res, err = _post_with_progress(base, "/ingest_linkedin", {"urls": urls, "cookie": cookie}, 12.0, "在线采集")
                if not err and res:
                    items = res.get("items", [])
                    jobs_pool = st.session_state.get("hr_jobs_pool") or []
                    for it in items:
                        jobs_pool.append({"id": str(it.get("url") or len(jobs_pool)), "text": str(it.get("text", ""))})
                    st.session_state["hr_jobs_pool"] = jobs_pool
                    st.success(f"已采集并合并 {len(items)} 条岗位描述")
        pool = st.session_state.get("hr_jobs_pool") or []
        if pool:
            opts = [p.get("id") for p in pool]
            sel = st.selectbox("岗位ID", options=opts, index=0, key="hr_job_select")
            chosen = next((x for x in pool if x.get("id") == sel), None)
            st.session_state["hr_chosen_job_text"] = chosen.get("text") if chosen else None
            st.markdown("**三级漏斗量化输出**")
            top_k = st.number_input("召回TopK", min_value=10, value=50, step=10, key="hr_funnel_topk")
            rules_text = st.text_area("自定义规则(JSON)", value="[]", height=80, key="hr_funnel_rules")
            if st.button("入库并解释（一步）", key="hr_do_funnel_index_explain"):
                items = []
                resumes = st.session_state.get("hr_resumes") or []
                for r in resumes:
                    items.append({"id": str(r.get("id")), "text": str(r.get("text", ""))})
                if items:
                    api_post(base, "/vector_index", {"items": items}, timeout=8.0)
                try:
                    rules = json.loads(rules_text) if rules_text.strip() else []
                except Exception:
                    rules = []
                payload = {"job_desc": st.session_state.get("hr_chosen_job_text") or "", "top_k": int(top_k), "custom_rules": rules}
                res, err = _post_with_progress(base, "/funnel_explain", payload, 12.0, "执行漏斗")
                if err:
                    st.error(err)
                else:
                    st.metric("语义召回数", str(res.get("recall_count", 0)))
                    st.metric("规则通过数", str(res.get("rule_passed_count", 0)))
                    passed = res.get("passed", [])
                    items = res.get("items", [])
                    if passed:
                        st.markdown("**通过候选（综合排序）**")
                        st.dataframe([
                            {
                                "id": it.get("id"),
                                "base": it.get("base"),
                                "skill": it.get("skill"),
                                "implicit": it.get("implicit"),
                                "format": it.get("format"),
                                "score": it.get("score"),
                                "matched_domains": it.get("matched_domains", []),
                                "matched_methods": it.get("matched_methods", []),
                                "matched_certs": it.get("matched_certs", []),
                                "matched_lang_levels": it.get("matched_lang_levels", []),
                                "proficiency_list": it.get("proficiency_list", []),
                                "proficiency_score": it.get("proficiency_score", 0.0),
                                "metric_score": it.get("metric_score", 0.0),
                                "explain": it.get("explain", "")
                            }
                        for it in passed
                        ], use_container_width=True)
                    fails = [it for it in items if not it.get("rule_ok")]
                    if fails:
                        st.markdown("**被过滤候选（原因）**")
                        st.dataframe([
                            {"id": it.get("id"), "reason": it.get("rule_reason"), "base": it.get("base")}
                            for it in fails
                        ], use_container_width=True)
                    st.markdown("**计算过程说明**")
                    st.info("综合分 = 0.6×基础相似度 + 0.2×技能覆盖比例 + 0.15×隐性需求匹配度 + 0.05×格式质量；学历/年限/职位为匹配细项（用于候选画像展示，不计入综合分）")
        else:
            st.session_state["hr_chosen_job_text"] = None
    if resumes:
        st.markdown("<div class='step-title'>步骤3：匹配与评分</div>", unsafe_allow_html=True)
        st.caption("基于JD与候选简历执行匹配并输出三级漏斗量化过程")
        rules_text_step3 = st.text_area("自定义规则(JSON)", value="[]", height=80, key="hr_match_rules")
        top_k_step3 = st.number_input("召回TopK", min_value=10, value=50, step=10, key="hr_match_topk")
        run = st.button("执行匹配并展示过程", key="hr_do_match")
        if run:
            jd_text_use = jd_text or ""
            if st.session_state.get("hr_struct_jd"):
                sj = st.session_state["hr_struct_jd"]
                jd_text_use = jd_text_use or ("\n".join((sj.get("title") or "", *sj.get("requirements", []))) )
            chosen_text = st.session_state.get("hr_chosen_job_text")
            if chosen_text:
                jd_text_use = chosen_text
            results = []
            for r in resumes:
                res, err = api_post(base, "/match", {"resume_text": r.get("text", ""), "job_text": jd_text_use}, timeout=float(st.session_state.get("api_timeout", 12.0)))
                if err:
                    continue
                d = res.get("details", {})
                w = st.session_state.get("hr_weights", {})
                s_adj = sum(float(d.get(k, 0)) * float(w.get(k, 0)) for k in w.keys())
                results.append({"id": r.get("id"), "score": res.get("score"), "score_adj": s_adj, "details": d, "resume_profile": res.get("resume_profile", {}), "job_profile": res.get("job_profile", {})})
            results = sorted(results, key=lambda x: x.get("score_adj", x.get("score", 0.0)), reverse=True)[:10]
            st.session_state["hr_match_results"] = results
            items = []
            for r in resumes:
                items.append({"id": str(r.get("id")), "text": str(r.get("text", ""))})
            if items:
                api_post(base, "/vector_index", {"items": items}, timeout=8.0)
            try:
                rules = json.loads(rules_text_step3) if rules_text_step3.strip() else []
            except Exception:
                rules = []
            payload = {"job_desc": jd_text_use, "top_k": int(top_k_step3), "custom_rules": rules}
            funnel, ferr = _post_with_progress(base, "/funnel_explain", payload, 12.0, "执行漏斗")
            st.session_state["hr_funnel_explain"] = funnel if not ferr else None
        rows = st.session_state.get("hr_match_results") or []
        if rows:
            st.dataframe([{ "id": x.get("id"), "score": x.get("score"), "score_adj": x.get("score_adj") } for x in rows], use_container_width=True)
            _download_excel_or_csv_sheets({"results": rows}, "hr_match_results.xlsx", "hr_match_results", prefer_excel=True)
            st.markdown("**三级漏斗量化过程（统一展示）**")
            fe = st.session_state.get("hr_funnel_explain") or {}
            if fe:
                st.metric("语义召回数", str(fe.get("recall_count", 0)))
                st.metric("规则通过数", str(fe.get("rule_passed_count", 0)))
                passed = fe.get("passed", [])
                # 合并匹配细项（学历/年限/职位）
                match_map = {str(x.get("id")): x for x in rows}
                merged = []
                for it in passed:
                    rid = str(it.get("id"))
                    det = match_map.get(rid, {}).get("details", {})
                    merged.append({
                        "id": rid,
                        "base": it.get("base"),
                        "skill": it.get("skill"),
                        "implicit": it.get("implicit"),
                        "format": it.get("format"),
                        "score": it.get("score"),
                        "degree": float(det.get("degree_score", 0)),
                        "years": float(det.get("years_score", 0)),
                        "position": float(det.get("position_score", 0)),
                        "matched_skills_count": len(det.get("matched_skills", [])),
                        "matched_skills": det.get("matched_skills", []),
                        "matched_domains": it.get("matched_domains", []),
                        "matched_methods": it.get("matched_methods", []),
                        "proficiency_score": float(it.get("proficiency_score", 0.0)),
                        "metric_score": float(it.get("metric_score", 0.0)),
                    })
                st.dataframe(merged, use_container_width=True)
                st.markdown("**计算过程说明**")
                st.info("综合分 = 0.6×基础相似度 + 0.2×技能覆盖比例 + 0.15×隐性需求匹配度 + 0.05×格式质量；学历/年限/职位为匹配细项（用于候选画像展示，不计入综合分）")
            # 补充图表：Top3雷达 + 能力占比饼图（Top3平均）
            if st_echarts and rows:
                c1, c2 = st.columns(2)
                with c1:
                    st.markdown("**Top3 综合雷达对比**")
                    top = rows[:3]
                    if top:
                        inds = [
                            ("技能", "skill_ratio"), ("学历", "degree_score"), ("年限", "years_score"), ("职位", "position_score"),
                            ("关键词", "keyword_ratio"), ("证书", "certs_score"), ("语言", "languages_ratio"), ("格式", "format_score")
                        ]
                        indicator = [{"name": n, "max": 1.0} for n, _ in inds]
                        radar_data = []
                        names = []
                        for it in top:
                            det = it.get("details", {})
                            vals = [float(det.get(key, 0)) for _, key in inds]
                            radar_data.append({"value": vals, "name": str(it.get("id"))})
                            names.append(str(it.get("id")))
                        r_opts = {"legend": {"data": names}, "radar": {"indicator": indicator}, "series": [{"type": "radar", "data": radar_data, "itemStyle": {"color": "#1976d2"}}]}
                        st_echarts(r_opts, height=360)
                with c2:
                    st.markdown("**Top3 平均能力结构占比**")
                    top = rows[:3]
                    if top:
                        hard_sum = 0.0
                        soft_sum = 0.0
                        fmt_sum = 0.0
                        for it in top:
                            d = it.get("details", {})
                            hard_sum += (float(d.get("degree_score", 0)) + float(d.get("years_score", 0)) + float(d.get("position_score", 0))) / 3.0
                            soft_sum += float(d.get("skill_ratio", 0))
                            fmt_sum += float(d.get("format_score", 0))
                        cnt = max(1, len(top))
                        pie_data = [
                            {"name": "硬性", "value": round(hard_sum / cnt, 4)},
                            {"name": "软性", "value": round(soft_sum / cnt, 4)},
                            {"name": "格式", "value": round(fmt_sum / cnt, 4)},
                        ]
                        p_opts = {"tooltip": {"trigger": "item"}, "series": [{"type": "pie", "radius": "55%", "data": pie_data}]}
                        st_echarts(p_opts, height=360)
            else:
                import pandas as pd
                top = rows[:3]
                if top:
                    inds = [
                        ("技能", "skill_ratio"), ("学历", "degree_score"), ("年限", "years_score"), ("职位", "position_score"),
                        ("关键词", "keyword_ratio"), ("证书", "certs_score"), ("语言", "languages_ratio"), ("格式", "format_score")
                    ]
                    data = {}
                    for it in top:
                        det = it.get("details", {})
                        data[str(it.get("id"))] = [float(det.get(key, 0)) for _, key in inds]
                    df_radar = pd.DataFrame(data, index=[n for n, _ in inds])
                    st.bar_chart(df_radar)
                    hard_sum = 0.0
                    soft_sum = 0.0
                    fmt_sum = 0.0
                    for it in top:
                        d = it.get("details", {})
                        hard_sum += (float(d.get("degree_score", 0)) + float(d.get("years_score", 0)) + float(d.get("position_score", 0))) / 3.0
                        soft_sum += float(d.get("skill_ratio", 0))
                        fmt_sum += float(d.get("format_score", 0))
                    cnt = max(1, len(top))
                    df_pie = pd.DataFrame({"value": [round(hard_sum / cnt, 4), round(soft_sum / cnt, 4), round(fmt_sum / cnt, 4)]}, index=["硬性", "软性", "格式"])
                    st.bar_chart(df_pie)
            
            # 候选人详情选择（Top10）
            st.markdown("**候选人详情（Top10）**")
            sel_ids = [str(it.get("id")) for it in rows[:10]]
            if sel_ids:
                sel = st.selectbox("查看候选人详情（Top10）", options=sel_ids, index=0, key="hr_top_sel")
                chosen = next((x for x in rows[:10] if str(x.get("id")) == sel), None)
                if chosen:
                    _render_candidate_insights(
                        chosen.get("details", {}),
                        chosen.get("resume_profile", {}),
                        chosen.get("job_profile", {}),
                        next((x.get("text") for x in resumes if x.get("id") == chosen.get("id")), ""),
                        jd_text,
                        base,
                    )
            
            
            
            # 步骤4：仅综合分析
            st.markdown("<div class='step-title'>步骤4：综合结果建议</div>", unsafe_allow_html=True)
            st.caption("输出综合分析（优势/风险）")
            # 先展示综合分析
            if sel:
                st.markdown("**综合分析（优势/风险）**")
                provider = st.session_state.get("llm_provider", "local")
                model = st.session_state.get("llm_model", "")
                rt = next((x for x in resumes if str(x.get("id")) == str(sel)), {})
                rep, er2 = api_post(base, "/evaluation_report", {"resume_text": rt.get("text", ""), "job_text": jd_text, "provider": provider, "model": model}, timeout=6.0)
                if er2:
                    st.error(er2)
                else:
                    txt = rep.get("report", "")
                    adv_txt = ""
                    risk_txt = ""
                    if txt:
                        import re
                        m_adv = re.search(r"优势[:：]", txt)
                        m_risk = re.search(r"风险[:：]", txt)
                        if m_adv and m_risk:
                            a_start = m_adv.end()
                            r_start = m_risk.end()
                            if a_start <= m_risk.start():
                                adv_txt = txt[a_start:m_risk.start()].strip()
                                risk_txt = txt[r_start:].strip()
                            else:
                                risk_txt = txt[r_start:m_adv.start()].strip()
                                adv_txt = txt[a_start:].strip()
                        elif m_adv:
                            a_start = m_adv.end()
                            seg = txt[a_start:]
                            adv_txt = re.split(r"[\n。；;]", seg)[0].strip()
                        elif m_risk:
                            r_start = m_risk.end()
                            seg = txt[r_start:]
                            risk_txt = re.split(r"[\n。；;]", seg)[0].strip()
                        else:
                            parts = [p.strip() for p in re.split(r"[\n。；;]", txt) if p.strip()]
                            adv_txt = parts[0] if parts else ""
                            risk_txt = parts[1] if len(parts) > 1 else ""
                        adv_txt = re.sub(r"风险[:：].*", "", adv_txt).strip()
                        risk_txt = re.sub(r"优势[:：].*", "", risk_txt).strip()
                    st.markdown("**优势：**")
                    st.write(adv_txt or "无")
                    st.write("---")
                    st.markdown("**风险：**")
                    st.write(risk_txt or "无")
            
            

def page_resume_optimize():
    st.subheader("简历优化")
    base = st.session_state.api_base
    f = st.file_uploader("上传简历文件或粘贴文本", type=["txt", "md", "docx", "pdf"]) 
    text = st.text_area("简历文本", height=200)
    run = st.button("生成优化建议")
    if run:
        if f is not None and not text.strip():
            text = _read_uploaded_text(f)
        if not text.strip():
            st.warning("请提供简历文本或上传文件")
            return
        res, err = api_post(base, "/resume_optimize", {"text": text}, timeout=6.0)
        if err:
            st.error(err)
        else:
            st.json(res)

def page_interview_training():
    st.subheader("面试训练")
    base = st.session_state.api_base
    jd = st.text_area("岗位描述（可选）", height=120, key="interview_training_job_desc")
    resume_text = st.text_area("候选人简历文本", height=200, key="interview_training_resume_text")
    run = st.button("生成面试题", key="interview_training_generate")
    if run:
        if not resume_text.strip():
            st.warning("请先输入简历文本")
            return
        payload = {"job_desc": jd, "resume_text": resume_text}
        res, err = api_post(base, "/interview_questions", payload, timeout=8.0)
        if err:
            st.error(err)
        else:
            qs = res.get("questions", [])
            for i, q in enumerate(qs):
                st.write(f"问{i+1}：{q}")

def page_jobseeker_wizard():
    st.subheader("求职者流程向导")
    base = st.session_state.api_base
    rf = st.file_uploader("上传简历文件（txt/md/docx/pdf）", type=["txt", "md", "docx", "pdf"], key="jobseeker_resume_file")
    rtext = st.text_area("或粘贴简历文本", height=160, key="jobseeker_resume_text")
    if st.button("解析简历", key="jobseeker_parse_resume"):
        if rf is not None and not rtext.strip():
            rtext = _read_uploaded_text(rf)
        st.session_state["jobseeker_text"] = rtext
        st.success("已解析简历文本")
    text = st.session_state.get("jobseeker_text", "")
    if text:
        opt, err = api_post(base, "/resume_optimize", {"text": text}, timeout=8.0)
        if not err and opt:
            st.json(opt)
        rec, er2 = api_post(base, "/recommend_jobs", {"resume_text": text, "top_k": 5}, timeout=12.0)
        items = rec.get("items", []) if (not er2 and rec) else []
        if items:
            st.dataframe([{ "id": it.get("id"), "score": it.get("score") } for it in items], use_container_width=True)
            for it in items[:3]:
                mr, me = api_post(base, "/match", {"resume_text": text, "job_text": it.get("text", "")}, timeout=8.0)
                if me:
                    continue
                _render_match_charts(mr.get("details", {}))
            qs, qe = api_post(base, "/interview_questions", {"job_desc": "\n".join([i.get("text", "") for i in items[:3]]), "resume_text": text}, timeout=8.0)
            if not qe and qs:
                for i, q in enumerate(qs.get("questions", [])):
                    st.write(f"问{i+1}：{q}")

def page_candidate_analysis():
    st.subheader("匹配度分析")
    base = st.session_state.api_base
    resume_text = st.text_area("简历文本", height=200, key="candidate_analysis_resume_text")
    jobs_dir = st.text_input("岗位目录", os.path.join("data", "raw_jobs"), key="candidate_analysis_jobs_dir")
    run = st.button("分析前三最匹配岗位")
    if run:
        if not resume_text.strip():
            st.warning("请先填写简历文本")
            return
        res, err = api_post(base, "/recommend_jobs", {"resume_text": resume_text, "jobs_dir": jobs_dir, "top_k": 3}, timeout=12.0)
        if err:
            st.error(err)
            return
        items = res.get("items", [])
        for it in items:
            jt = it.get("file") or it.get("id")
            st.info(f"岗位：{jt}")
            mr, me = api_post(base, "/match", {"resume_text": resume_text, "job_text": it.get("text", "")})
            if me:
                continue
            _render_match_charts(mr.get("details", {}))
        gen = st.button("生成AI评估报告并下载", key="candidate_analysis_generate_report")
        if gen:
            rep, err = api_post(base, "/evaluation_report", {"resume_text": resume_text, "job_text": "\n".join([i.get("text", "") for i in items])})
            if err:
                st.error(err)
            else:
                txt = rep.get("report", "")
                st.text_area("评估报告", value=txt, height=200, key="candidate_analysis_report_text")
                st.download_button("下载评估报告TXT", data=txt.encode("utf-8"), file_name="candidate_evaluation.txt")
                html = f"""
                <html><head><meta charset='utf-8'><title>评估报告</title></head><body>
                <h2>岗位Top3匹配度分析</h2>
                <p>{txt}</p>
                <h3>优势</h3>
                <ul>
                <li>技能命中：基于匹配细项的技能交集</li>
                <li>学历/年限/职位关键词满足度</li>
                </ul>
                <h3>不足与建议</h3>
                <ul>
                <li>补充缺失的岗位核心技能案例与量化成果</li>
                <li>完善学历/年限表述与验证方式</li>
                </ul>
                </body></html>
                """
                st.download_button("下载评估报告HTML", data=html.encode("utf-8"), file_name="candidate_evaluation.html")

def page_job_recommend():
    st.subheader("岗位推荐")
    base = st.session_state.api_base
    resume_text = st.text_area("候选人简历文本", height=200)
    src = st.radio("数据源", ["目录", "SQLite"], index=0)
    jobs_dir = (
        st.text_input("岗位目录", os.path.join("data", "raw_jobs"), key="job_recommend_jobs_dir")
        if src == "目录"
        else st.text_input("SQLite路径", value=os.getenv("RECOMMEND_JOBS_SQLITE", ""), key="job_recommend_sqlite_path")
    )
    top_k = st.number_input("Top N", min_value=1, value=5, step=1)
    fusion = st.slider("融合权重(匹配分α vs 语义相似度1-α)", 0.0, 1.0, 0.7, 0.05)
    colf1, colf2, colf3 = st.columns(3)
    with colf1:
        industry = st.text_input("行业关键字", value="")
    with colf2:
        region = st.text_input("地区关键字", value="")
    with colf3:
        smin = st.number_input("最低薪资(￥)", min_value=0.0, value=0.0, step=1000.0)
    smax = st.number_input("最高薪资(￥)", min_value=0.0, value=0.0, step=1000.0)
    run = st.button("生成推荐")
    if run:
        if not resume_text.strip():
            st.warning("请先填写简历文本")
            return
        payload = {"resume_text": resume_text, "top_k": int(top_k), "industry": industry, "region": region}
        if smin > 0:
            payload["salary_min"] = float(smin)
        if smax > 0:
            payload["salary_max"] = float(smax)
        if src == "目录":
            payload["jobs_dir"] = jobs_dir
            payload["data_source"] = "dir"
        else:
            payload["sqlite_path"] = jobs_dir
            payload["data_source"] = "sqlite"
        payload["fusion_weight"] = float(fusion)
        res, err = api_post(base, "/recommend_jobs", payload, timeout=12.0)
        if err:
            st.error(err)
        else:
            items = res.get("items", [])
            rows = [{"id": it.get("id"), "score": it.get("score") } for it in items]
            st.dataframe(rows, use_container_width=True)
            _download_excel_or_csv_sheets({"recommended": rows}, "recommended_jobs.xlsx", "recommended_jobs", prefer_excel=True)

def page_admin_monitor():
    st.subheader("流程监控")
    log_train = os.path.join("logs", "training_metrics.jsonl")
    log_eval = os.path.join("logs", "ner_eval_metrics.jsonl")
    tabs = st.tabs(["训练指标", "评估指标", "特征提取"]) 
    with tabs[0]:
        if os.path.isfile(log_train):
            import json
            xs, loss = [], []
            prec, rec, f1 = [], [], []
            with open(log_train, "r", encoding="utf-8") as f:
                for i, line in enumerate(f):
                    try:
                        obj = json.loads(line)
                        xs.append(obj.get("step", i))
                        loss.append(float(obj.get("loss", 0)))
                        prec.append(float(obj.get("precision", 0)))
                        rec.append(float(obj.get("recall", 0)))
                        f1.append(float(obj.get("f1", 0)))
                    except Exception:
                        pass
            try:
                if st_echarts:
                    opts = {
                        "tooltip": {"trigger": "axis"},
                        "xAxis": {"type": "category", "data": xs},
                        "yAxis": {"type": "value"},
                        "series": [
                            {"type": "line", "name": "loss", "data": loss},
                            {"type": "line", "name": "precision", "data": prec},
                            {"type": "line", "name": "recall", "data": rec},
                            {"type": "line", "name": "f1", "data": f1},
                        ]
                    }
                    st_echarts(opts, height=320)
                else:
                    import pandas as pd
                    df = pd.DataFrame({"step": xs, "loss": loss, "precision": prec, "recall": rec, "f1": f1})
                    st.line_chart(df.set_index("step"))
            except Exception:
                st.error("训练日志解析失败")
        else:
            st.info("未检测到训练日志：logs/training_metrics.jsonl")
    with tabs[1]:
        if os.path.isfile(log_eval):
            import json
            xs, prec, rec, f1 = [], [], [], []
            with open(log_eval, "r", encoding="utf-8") as f:
                for i, line in enumerate(f):
                    try:
                        obj = json.loads(line)
                        xs.append(obj.get("step", i))
                        prec.append(float(obj.get("precision", 0)))
                        rec.append(float(obj.get("recall", 0)))
                        f1.append(float(obj.get("f1", 0)))
                    except Exception:
                        pass
            try:
                if st_echarts:
                    opts = {
                        "tooltip": {"trigger": "axis"},
                        "xAxis": {"type": "category", "data": xs},
                        "yAxis": {"type": "value"},
                        "series": [
                            {"type": "line", "name": "precision", "data": prec},
                            {"type": "line", "name": "recall", "data": rec},
                            {"type": "line", "name": "f1", "data": f1},
                        ]
                    }
                    st_echarts(opts, height=320)
                else:
                    import pandas as pd
                    df = pd.DataFrame({"step": xs, "precision": prec, "recall": rec, "f1": f1})
                    st.line_chart(df.set_index("step"))
            except Exception:
                st.error("评估日志解析失败")
        else:
            st.info("未检测到评估日志：logs/ner_eval_metrics.jsonl")
    with tabs[2]:
        import json
        data_path = os.path.join("data", "processed", "resumes_for_annotation.json")
        if not os.path.isfile(data_path):
            st.info("未检测到处理后的简历数据集：data/processed/resumes_for_annotation.json")
        else:
            try:
                with open(data_path, "r", encoding="utf-8") as f:
                    ds = json.load(f)
                texts = [d.get("text", "") for d in ds]
                total = max(len(texts), 1)
                import re
                degree_cnt = {}
                years_vals = []
                skills_cnt = {}
                langs_cnt = {}
                fw_cnt = {}
                tools_cnt = {}
                city_cnt = {}
                langs = {"java","python","go","rust","c++","c#","js","ts","sql"}
                fws = {"spring","springboot","django","flask","fastapi","vue","react","angular","spark","hadoop"}
                tools = {"docker","k8s","kubernetes","git","jenkins","maven","gradle","terraform"}
                cities = ["北京","上海","广州","深圳","杭州","南京","成都","武汉","西安","苏州","天津","重庆","合肥","厦门","沈阳","大连","无锡","佛山","宁波","青岛"]
                for t in texts:
                    deg_m = re.findall(r"博士|硕士|本科|大专", t)
                    if deg_m:
                        degree_cnt[deg_m[0]] = degree_cnt.get(deg_m[0], 0) + 1
                    ym = re.findall(r"(\d+)\s*年", t)
                    if ym:
                        try:
                            years_vals.append(int(ym[0]))
                        except Exception:
                            pass
                    toks = [x.lower() for x in re.findall(r"[A-Za-z][A-Za-z0-9+#\.\-]{1,}", t)]
                    for tok in set(toks):
                        skills_cnt[tok] = skills_cnt.get(tok, 0) + 1
                        if tok in langs:
                            langs_cnt[tok] = langs_cnt.get(tok, 0) + 1
                        if tok in fws:
                            fw_cnt[tok] = fw_cnt.get(tok, 0) + 1
                        if tok in tools:
                            tools_cnt[tok] = tools_cnt.get(tok, 0) + 1
                    for c in cities:
                        if c in t:
                            city_cnt[c] = city_cnt.get(c, 0) + 1
                import pandas as pd, numpy as np
                deg_cov = sum(degree_cnt.values()) * 100.0 / total
                years_cov = (len(years_vals) * 100.0 / total)
                avg_chars = sum(len(x) for x in texts) / float(total)
                avg_words = sum(len(re.findall(r"\w+", x)) for x in texts) / float(total)
                avg_skills = sum(len(set(re.findall(r"[A-Za-z][A-Za-z0-9+#\.\-]{1,}", x.lower()))) for x in texts) / float(total)
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("样本数", f"{total}")
                c2.metric("学历覆盖率(%)", f"{deg_cov:.1f}")
                c3.metric("年限覆盖率(%)", f"{years_cov:.1f}")
                c4.metric("平均技能词数", f"{avg_skills:.1f}")
                df_deg = pd.DataFrame({"degree": list(degree_cnt.keys()), "count": list(degree_cnt.values())})
                if not df_deg.empty:
                    st.write("学历分布")
                    st.bar_chart(df_deg.set_index("degree"))
                top_sk = sorted(skills_cnt.items(), key=lambda x: x[1], reverse=True)[:30]
                if top_sk:
                    df_sk = pd.DataFrame(top_sk, columns=["skill","count"])
                    st.write("Top30技能词分布")
                    st.bar_chart(df_sk.set_index("skill"))
                if langs_cnt:
                    df_lang = pd.DataFrame(sorted(langs_cnt.items(), key=lambda x: x[1], reverse=True), columns=["language","count"])
                    st.write("编程语言分布")
                    st.bar_chart(df_lang.set_index("language"))
                if fw_cnt:
                    df_fw = pd.DataFrame(sorted(fw_cnt.items(), key=lambda x: x[1], reverse=True), columns=["framework","count"])
                    st.write("框架分布")
                    st.bar_chart(df_fw.set_index("framework"))
                if tools_cnt:
                    df_tools = pd.DataFrame(sorted(tools_cnt.items(), key=lambda x: x[1], reverse=True), columns=["tool","count"])
                    st.write("工具分布")
                    st.bar_chart(df_tools.set_index("tool"))
                if years_vals:
                    hist, bins = np.histogram(years_vals, bins=min(10, max(3, len(set(years_vals)))))
                    centers = [f"{int(bins[i])}-{int(bins[i+1])}" for i in range(len(bins)-1)]
                    df_hist = pd.DataFrame({"bucket": centers, "count": hist})
                    st.write("年限直方图")
                    st.bar_chart(df_hist.set_index("bucket"))
                if city_cnt:
                    df_city = pd.DataFrame(sorted(city_cnt.items(), key=lambda x: x[1], reverse=True), columns=["city","count"])
                    st.write("城市提及分布")
                    st.bar_chart(df_city.set_index("city"))
                st.success("统计完成")
            except Exception:
                st.error("画像统计失败")

def _can_upload() -> bool:
    perms = st.session_state.get("permissions") or []
    return "uploads.write" in perms

def page_interview_feedback():
    st.subheader("评价记录")
    if "feedback_rows" not in st.session_state:
        st.session_state["feedback_rows"] = []
    cid = st.text_input("候选人ID")
    score = st.number_input("评分(0-100)", min_value=0, max_value=100, value=80, step=1)
    notes = st.text_area("评价备注", height=160)
    add = st.button("添加记录")
    if add:
        st.session_state["feedback_rows"].append({"id": cid, "score": int(score), "notes": notes, "ts": int(time.time())})
        st.success("已添加")
    rows = st.session_state.get("feedback_rows", [])
    if rows:
        import pandas as pd, json
        df = pd.DataFrame(rows)
        st.dataframe(df, use_container_width=True)
        st.download_button("下载评价CSV", data=df.to_csv(index=False).encode("utf-8-sig"), file_name="interview_feedback.csv")
        save_srv = st.button("保存到后端")
        if save_srv:
            if _can_upload():
                try:
                    payload = json.dumps(rows, ensure_ascii=False, indent=2)
                    res, err = api_post(st.session_state.api_base, "/uploads", {"text": payload, "filename": "interview_feedback.json"})
                    if err:
                        st.error(f"上传失败：{err}")
                    else:
                        st.success("已保存到后端")
                except Exception as e:
                    st.error(str(e))
            else:
                st.warning("当前令牌无上传权限，已提供本地下载功能")


 
def page_funnel():
    st.subheader("三级漏斗筛选")
    base = st.session_state.api_base
    job_desc = st.text_area("岗位描述", height=160)
    source = st.radio("候选来源", ["使用处理后的JSON", "向量库已有"], index=0)
    top_k = st.number_input("Top K", min_value=1, value=50, step=1, key="funnel_top_k")
    rules_text = st.text_area("自定义规则(JSON)", value='[{"field":"years","operator":"gt","value":2}]', height=100, key="funnel_rules")
    if st.button("执行漏斗"):
        try:
            rules = json.loads(rules_text) if rules_text.strip() else []
        except Exception:
            st.error("规则JSON不合法")
            return
        if not job_desc.strip():
            st.warning("请填写岗位描述")
            return
        if source == "使用处理后的JSON":
            resume_json = os.path.join("data", "processed", "resumes_for_annotation.json")
            if not os.path.isfile(resume_json):
                st.error("简历JSON不存在")
                return
            with open(resume_json, "r", encoding="utf-8") as f:
                data = json.load(f)
            items = []
            for i, d in enumerate(data):
                items.append({"id": str(i), "text": d.get("text", "")})
            _post_with_progress(base, "/vector_index", {"items": items}, 8.0, "向量入库")
        res, err = _post_with_progress(base, "/filter", {"job_desc": job_desc, "top_k": int(top_k), "custom_rules": rules}, 12.0, "执行漏斗")
        if err:
            st.error(f"调用失败：{err}")
            return
        results = res.get("results", [])
        st.success(f"完成筛选，返回 {len(results)} 条")
        if results:
            rows = [{"id": r.get("id"), "score": r.get("score"), "base": r.get("base"), "skill": r.get("skill"), "implicit": r.get("implicit"), "format": r.get("format") } for r in results]
            st.dataframe(rows, use_container_width=True)
            if st_echarts:
                xs = [r.get("id") for r in results[:10]]
                ys = [r.get("score") for r in results[:10]]
                opts = {
                    "tooltip": {"trigger": "axis"},
                    "xAxis": {"type": "category", "data": xs},
                    "yAxis": {"type": "value", "name": "score"},
            "series": [{"type": "bar", "data": ys, "itemStyle": {"color": "#1976d2"}}]
                }
                st_echarts(opts, height=320)
            ids_all = [str(r.get("id")) for r in results]
            default_sel = ids_all[:5]
            sel = st.multiselect("选择对比候选人ID", ids_all[:20], default=default_sel)
            if st_echarts:
                inds = ["base", "skill", "implicit", "format"]
                chosen = [r for r in results if str(r.get("id")) in sel][:20]
                maxs = [1.0, 1.0, 1.0, 1.0]
                if chosen:
                    maxs = [max(0.1, max(float(x.get("base",0)) for x in chosen)), max(0.1, max(float(x.get("skill",0)) for x in chosen)), max(0.1, max(float(x.get("implicit",0)) for x in chosen)), max(0.1, max(float(x.get("format",0)) for x in chosen))]
                ind_cfg = [{"name": n, "max": float(m)} for n, m in zip(inds, maxs)]
                radar_data = []
                for r in chosen:
                    radar_data.append({"value": [float(r.get("base",0)), float(r.get("skill",0)), float(r.get("implicit",0)), float(r.get("format",0))], "name": str(r.get("id"))})
                r_opts = {"legend": {"data": [str(r.get("id")) for r in chosen]}, "radar": {"indicator": ind_cfg}, "series": [{"type": "radar", "data": radar_data, "itemStyle": {"color": "#1976d2"}}]}
                st_echarts(r_opts, height=360)
            import pandas as pd
            df = pd.DataFrame(rows)
            csv = df.to_csv(index=False).encode("utf-8-sig")
            st.download_button("下载CSV", data=csv, file_name="funnel_results.csv")

def page_fairness():
    st.subheader("公平性报告")
    dev_thr = st.slider("偏差阈值(%)", min_value=1, max_value=10, value=2, step=1)
    resume_json = os.path.join("data", "processed", "resumes_for_annotation.json")
    if not os.path.isfile(resume_json):
        st.info("未找到预处理输出")
        return
    with open(resume_json, "r", encoding="utf-8") as f:
        data = json.load(f)
    gender_cnt = {}
    tier_cnt = {}
    gap_cnt = {}
    for d in data:
        tags = d.get("fairness_tags", {})
        g = tags.get("gender", "未知")
        t = tags.get("school_tier", "未知")
        gp = tags.get("gap_type", "未知")
        gender_cnt[g] = gender_cnt.get(g, 0) + 1
        tier_cnt[t] = tier_cnt.get(t, 0) + 1
        gap_cnt[gp] = gap_cnt.get(gp, 0) + 1
    # 计算百分比
    def _to_percent(d):
        s = max(1, sum(d.values()))
        return {k: round(v * 100.0 / s, 2) for k, v in d.items()}
    g_pct = _to_percent(gender_cnt)
    t_pct = _to_percent(tier_cnt)
    gp_pct = _to_percent(gap_cnt)
    if st_echarts:
        opts_g = {
            "tooltip": {"trigger": "axis"},
            "xAxis": {"type": "category", "data": list(g_pct.keys())},
            "yAxis": {"type": "value", "name": "通过率(%)"},
            "series": [{"type": "bar", "data": list(g_pct.values()), "itemStyle": {"color": "#1976d2"}}],
            "markLine": {"data": [{"yAxis": dev_thr}]}
        }
        st_echarts(opts_g, height=300)
        opts_t = {
            "tooltip": {"trigger": "axis"},
            "xAxis": {"type": "category", "data": list(t_pct.keys())},
            "yAxis": {"type": "value", "name": "通过率(%)"},
            "series": [{"type": "bar", "data": list(t_pct.values()), "itemStyle": {"color": "#1976d2"}}],
            "markLine": {"data": [{"yAxis": dev_thr}]}
        }
        st_echarts(opts_t, height=300)
        opts_gp = {
            "tooltip": {"trigger": "axis"},
            "xAxis": {"type": "category", "data": list(gp_pct.keys())},
            "yAxis": {"type": "value", "name": "通过率(%)"},
            "series": [{"type": "bar", "data": list(gp_pct.values()), "itemStyle": {"color": "#1976d2"}}],
            "markLine": {"data": [{"yAxis": dev_thr}]}
        }
        st_echarts(opts_gp, height=300)
    else:
        st.json({"gender": g_pct, "school_tier": t_pct, "gap_type": gp_pct})
    # 导出CSV
    import pandas as pd
    df = pd.DataFrame({
        "gender": g_pct,
        "school_tier": t_pct,
        "gap_type": gp_pct,
    })
    csv = df.to_csv().encode("utf-8-sig")
    st.download_button("下载公平性百分比CSV", data=csv, file_name="fairness_report.csv")
def page_decision():
    st.subheader("决策辅助")
    base = st.session_state.api_base
    job_desc = st.text_area("岗位描述（可一行一个支持批量）", height=160)
    rules_text = st.text_area("自定义规则(JSON)", value='[{"field":"years","operator":"gt","value":2}]', height=100, key="decision_rules")
    top_k = st.number_input("Top K", min_value=1, value=50, step=1, key="decision_top_k")
    page = st.number_input("页码", min_value=1, value=1, step=1)
    page_size = st.number_input("每页数量", min_value=5, value=20, step=5)
    llm_flag = st.checkbox("启用大模型面试题", value=bool(st.session_state.get("llm_global_enabled", False)), key="decision_llm")
    csv_flag = st.checkbox("导出为CSV", value=bool(st.session_state.get("export_csv_default", False)), key="decision_export_csv")
    if st.button("应用面试题生成设置", key="decision_llm_apply"):
        _ = api_post(base, "/config/llm_enabled", {"enabled": bool(llm_flag)})
        st.success("已应用设置")
    run = st.button("生成推荐与分数线")
    if run:
        try:
            rules = json.loads(rules_text) if rules_text.strip() else []
        except Exception:
            st.error("规则JSON不合法")
            return
        if not job_desc.strip():
            st.warning("请填写岗位描述")
            return
        lines = [l.strip() for l in job_desc.split("\n") if l.strip()]
        if len(lines) > 1:
            payload = {"job_descs": lines, "top_k": int(top_k), "custom_rules": rules}
            res, err = _post_with_progress(base, "/decision_batch", payload, 20.0, "生成推荐与分数线")
            if err:
                st.error(f"调用失败：{err}")
                return
            items = res.get("items", [])
            for it in items:
                st.markdown("---")
                st.write("岗位：", it.get("job_desc"))
                decision = it.get("decision", {})
                threshold = decision.get("threshold")
                picks = decision.get("recommended", [])
                st.metric("动态分数线", f"{threshold}")
                st.dataframe([{"id": p.get("id"), "score": p.get("score")} for p in picks], use_container_width=True)
            overview_rows = []
            explain_rows = []
            for it in items:
                decision = it.get("decision", {})
                threshold = decision.get("threshold")
                picks = decision.get("recommended", [])
                overview_rows.append({"job_desc": it.get("job_desc"), "threshold": threshold, "recommended_count": len(picks)})
                try:
                    fe, fe_err = api_post(base, "/funnel_explain", {"job_desc": it.get("job_desc"), "top_k": int(top_k), "custom_rules": rules}, timeout=12.0)
                    passed = fe.get("passed", []) if fe and not fe_err else []
                    mm = {str(x.get("id")): x for x in passed}
                    for p in picks:
                        rid = str(p.get("id"))
                        ex = mm.get(rid, {})
                        explain_rows.append({
                            "job_desc": it.get("job_desc"),
                            "id": rid,
                            "score": p.get("score"),
                            "base": ex.get("base"),
                            "skill": ex.get("skill"),
                            "implicit": ex.get("implicit"),
                            "format": ex.get("format"),
                            "matched_domains": ex.get("matched_domains", []),
                            "matched_methods": ex.get("matched_methods", []),
                            "matched_certs": ex.get("matched_certs", []),
                            "matched_lang_levels": ex.get("matched_lang_levels", []),
                            "proficiency_list": ex.get("proficiency_list", []),
                            "proficiency_score": ex.get("proficiency_score", 0.0),
                            "metric_score": ex.get("metric_score", 0.0),
                            "explain": ex.get("explain", "")
                        })
                except Exception:
                    pass
            sheets = {"overview": overview_rows}
            if explain_rows:
                sheets["explain"] = explain_rows
            _download_excel_or_csv_sheets(sheets, "decision_batch_results.xlsx", "decision_batch", prefer_excel=(not csv_flag))
            return
        payload = {"job_desc": job_desc, "top_k": int(top_k), "custom_rules": rules, "page": int(page), "page_size": int(page_size)}
        res, err = _post_with_progress(base, "/decision", payload, 20.0, "生成推荐与分数线")
        if err:
            st.error(f"调用失败：{err}")
            return
        decision = res.get("decision", {})
        threshold = decision.get("threshold")
        picks = decision.get("recommended", [])
        st.metric("动态分数线", f"{threshold}")
        _summary_cards([("推荐数量", str(len(picks))), ("分数线", f"{threshold}"), ("Top1分数", f"{picks[0].get('score') if picks else 0}")])
        st.write("推荐名单：")
        st.dataframe([{"id": p.get("id"), "score": p.get("score")} for p in picks], use_container_width=True)
        st.write("分页结果：")
        page_rows = res.get("page_results", [])
        st.dataframe(page_rows, use_container_width=True)
        _download_excel_or_csv_sheets({
            "results": res.get("results", []),
            "recommended": [{"id": p.get("id"), "score": p.get("score"), "questions": " | ".join(p.get("interview_questions", []))} for p in picks]
        }, "decision_results.xlsx", "decision", prefer_excel=(not csv_flag))
        if picks and st_echarts:
            xs = [p.get("id") for p in picks]
            ys = [p.get("score") for p in picks]
            opts = {
                "tooltip": {"trigger": "axis"},
                "xAxis": {"type": "category", "data": xs},
                "yAxis": {"type": "value"},
                "series": [{"type": "line", "data": ys, "smooth": True, "itemStyle": {"color": "#1976d2"}}],
                "markLine": {"data": [{"yAxis": threshold}]}
            }
            st_echarts(opts, height=320)
        jrows = []
        for p in picks:
            jrows.append({"id": p.get("id"), "questions": " | ".join(p.get("interview_questions", []))})
        st.dataframe(jrows, use_container_width=True)
        import pandas as pd
        df = pd.DataFrame(res.get("results", []))
        csv = df.to_csv(index=False).encode("utf-8-sig")
        st.download_button("下载全量结果CSV", data=csv, file_name="decision_all_results.csv")

def page_jobseeker_wizard():
    base = st.session_state.api_base
    st.markdown("<div class='step-title'>步骤1：上传简历</div>", unsafe_allow_html=True)
    st.caption("上传文件或粘贴文本后解析，作为后续推荐与匹配的基础")
    rf = st.file_uploader("上传简历文件（txt/md/docx/pdf）", type=["txt", "md", "docx", "pdf"], key="jobseeker_resume_file")
    rtext = st.text_area("或粘贴简历文本", height=160, key="jobseeker_resume_text")
    if st.button("解析简历", key="jobseeker_parse_resume"):
        if rf is not None and not rtext.strip():
            rtext = _read_uploaded_text(rf)
        st.session_state["jobseeker_text"] = rtext
        st.success("已解析简历文本")
    text = st.session_state.get("jobseeker_text", "")
    if text:
        st.markdown("<div class='step-title'>步骤2：AI解析与优化</div>", unsafe_allow_html=True)
        st.caption("调用后端生成简历优化建议，提升岗位匹配度")
        opt, err = api_post(base, "/resume_optimize", {"text": text}, timeout=8.0)
        if not err and opt:
            st.json(opt)
        st.markdown("<div class='step-title'>步骤3：自动岗位推荐</div>", unsafe_allow_html=True)
        st.caption("根据简历画像推荐岗位并展示评分，可作为匹配候选")
        rec, er2 = api_post(base, "/recommend_jobs", {"resume_text": text, "top_k": 5}, timeout=12.0)
        if not er2 and rec:
            items = rec.get("items", [])
            st.dataframe([{ "id": it.get("id"), "score": it.get("score") } for it in items], use_container_width=True)
            st.markdown("<div class='step-title'>步骤4：综合匹配度（Top3）</div>", unsafe_allow_html=True)
            st.caption("展示Top3综合雷达对比，选择候选人查看详细匹配与建议")
            # 综合雷达对比
            top_items = items[:3]
            radar_ready = []
            for it in top_items:
                mr, me = api_post(base, "/match", {"resume_text": text, "job_text": it.get("text", "")}, timeout=8.0)
                if me:
                    continue
                radar_ready.append({"id": it.get("id"), "match": mr})
            if st_echarts and radar_ready:
                inds = [
                    ("技能", "skill_ratio"), ("学历", "degree_score"), ("年限", "years_score"), ("职位", "position_score"),
                    ("关键词", "keyword_ratio"), ("证书", "certs_score"), ("语言", "languages_ratio"), ("格式", "format_score")
                ]
                indicator = [{"name": n, "max": 1.0} for n, _ in inds]
                radar_data = []
                names = []
                for rr in radar_ready:
                    det = (rr.get("match") or {}).get("details", {})
                    vals = [float(det.get(key, 0)) for _, key in inds]
                    radar_data.append({"value": vals, "name": str(rr.get("id"))})
                    names.append(str(rr.get("id")))
                r_opts = {"legend": {"data": names}, "radar": {"indicator": indicator}, "series": [{"type": "radar", "data": radar_data, "itemStyle": {"color": "#1976d2"}}]}
                st_echarts(r_opts, height=360)
            # 单人详情选择
            sel_ids = [str(rr.get("id")) for rr in radar_ready]
            if sel_ids:
                sel = st.selectbox("查看候选人详情", options=sel_ids, index=0, key="jobseeker_top_sel")
                chosen = next((rr for rr in radar_ready if str(rr.get("id")) == sel), None)
                if chosen:
                    _render_candidate_insights(
                        (chosen.get("match") or {}).get("details", {}),
                        (chosen.get("match") or {}).get("resume_profile", {}),
                        (chosen.get("match") or {}).get("job_profile", {}),
                        text,
                        next((it.get("text", "") for it in top_items if str(it.get("id")) == sel), ""),
                        base,
                    )
            st.markdown("<div class='step-title'>步骤5：面试准备建议</div>", unsafe_allow_html=True)
            st.caption("基于匹配结果与简历内容生成面试题，辅助备考")
            qs, qe = api_post(base, "/interview_questions", {"job_desc": "\n".join([i.get("text", "") for i in items[:3]]), "resume_text": text}, timeout=8.0)
            if not qe and qs:
                for i, q in enumerate(qs.get("questions", [])):
                    st.write(f"问{i+1}：{q}")
def page_admin_wizard():
    base = st.session_state.api_base
    processed_path = os.path.join("data", "processed", "resumes_for_annotation.json")
    st.markdown("<div class='step-title'>步骤1：数据预处理</div>", unsafe_allow_html=True)
    st.caption("运行预处理生成标准化样本，输出前10条摘要便于核查")
    run_pp = st.button("运行预处理", key="admin_run_preprocess")
    if run_pp:
        try:
            if _preprocess_mod is not None and hasattr(_preprocess_mod, "process_resumes"):
                _preprocess_mod.process_resumes()
                st.success("预处理完成")
            else:
                st.warning("预处理模块不可用")
        except Exception as e:
            st.error(f"预处理失败：{e}")
    if os.path.isfile(processed_path):
        with open(processed_path, "r", encoding="utf-8") as f:
            try:
                ds = json.load(f)
                st.info(f"预处理样本数：{len(ds)}")
                rows = []
                for i, d in enumerate(ds[:10]):
                    rows.append({
                        "idx": i,
                        "id": d.get("id", i),
                        "text_len": len(d.get("text", "")),
                        "preview": (d.get("text", "")[:300] or ""),
                        "entities": len(d.get("entities", []) or [])
                    })
                if rows:
                    import pandas as pd
                    st.dataframe(pd.DataFrame(rows), use_container_width=True)
            except Exception:
                st.warning("预处理输出读取失败")

    st.markdown("<div class='step-title'>步骤2：模型特征提取</div>", unsafe_allow_html=True)
    st.caption("统计画像分布与覆盖率，生成基础特征图表以便监控")
    show_feats = st.button("计算画像分布", key="admin_show_features")
    if show_feats and os.path.isfile(processed_path):
        try:
            with open(processed_path, "r", encoding="utf-8") as f:
                ds = json.load(f)
            import re, pandas as pd
            degree_cnt = {}
            skills_cnt = {}
            years_vals = []
            for d in ds:
                text = d.get("text", "")
                deg_m = re.findall(r"博士|硕士|本科|大专", text)
                if deg_m:
                    degree_cnt[deg_m[0]] = degree_cnt.get(deg_m[0], 0) + 1
                ym = re.findall(r"(\d+)\s*年", text)
                if ym:
                    try:
                        years_vals.append(int(ym[0]))
                    except Exception:
                        pass
                sk = list({s.lower() for s in re.findall(r"[A-Za-z+#\.\-]{2,}", text)})
                for s in sk:
                    skills_cnt[s] = skills_cnt.get(s, 0) + 1
            total = max(len(ds), 1)
            st.metric("学历覆盖率(%)", f"{(sum(degree_cnt.values())*100.0/total):.1f}")
            st.metric("年限覆盖率(%)", f"{(len(years_vals)*100.0/total):.1f}")
            df_deg = pd.DataFrame({"degree": list(degree_cnt.keys()), "count": list(degree_cnt.values())})
            df_sk = pd.DataFrame(sorted(skills_cnt.items(), key=lambda x: x[1], reverse=True)[:20], columns=["skill", "count"])
            st.bar_chart(df_deg.set_index("degree"))
            st.bar_chart(df_sk.set_index("skill"))
            try:
                import matplotlib.pyplot as plt
                fig, ax = plt.subplots(figsize=(5,3))
                if years_vals:
                    ax.boxplot(years_vals, vert=True)
                ax.set_title("年限箱线图")
                st.pyplot(fig, use_container_width=False)
            except Exception:
                pass
        except Exception as e:
            st.error(f"特征提取失败：{e}")

    st.markdown("<div class='step-title'>步骤3：模型特性向量化</div>", unsafe_allow_html=True)
    st.caption("将简历文本向量入库，供后续检索与匹配使用")
    upsert = st.button("向量入库（简历）", key="admin_vector_upsert")
    if upsert:
        if not os.path.isfile(processed_path):
            st.error("未检测到预处理输出")
        else:
            with open(processed_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            items = [{"id": str(i), "text": d.get("text", "")} for i, d in enumerate(data)]
            res, err = _post_with_progress(base, "/vector_index", {"items": items}, 10.0, "向量入库")
            if err:
                st.error(err)
            else:
                st.success(f"入库完成：{res.get('count')} 条")

    st.markdown("<div class='step-title'>步骤4：实体模型生成过程（训练）</div>", unsafe_allow_html=True)
    st.caption("检测已训练模型，选择重训或直接开始训练并实时记录日志")
    model_path = os.path.join("models", "bert_entity", "model.safetensors")
    if os.path.isfile(model_path):
        st.success(f"检测到已训练模型：{model_path}")
        retrain = st.checkbox("重新训练（覆盖现有模型）", value=False, key="admin_retrain_flag")
        start_train = st.button("开始重新训练", key="admin_start_training_retrain") if retrain else None
    else:
        st.info("未检测到已训练模型，可一键启动训练，日志实时写入 logs/training_metrics.jsonl")
        start_train = st.button("开始训练实体模型", key="admin_start_training")
    if start_train:
        try:
            import subprocess, sys
            st.write("训练进度：")
            prog = st.progress(0)
            logbox = st.empty()
            # 通过 -u 保证 stdout 行刷出
            env = dict(os.environ)
            env["PYTHONIOENCODING"] = env.get("PYTHONIOENCODING", "utf-8")
            p = subprocess.Popen([sys.executable, "-u", "scripts/03_entity_model.py"], stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, env=env)
            lines = []
            steps = 0
            while True:
                line = p.stdout.readline()
                if not line:
                    if p.poll() is not None:
                        break
                    continue
                lines.append(line.rstrip())
                steps += 1
                logbox.text("\n".join(lines[-10:]))
                prog.progress(min(steps/100.0, 1.0))
            rc = p.poll()
            if rc == 0:
                st.success("训练完成")
            else:
                st.warning(f"训练进程退出码：{rc}")
        except Exception as e:
            st.error(f"启动训练失败：{e}")

    st.markdown("<div class='step-title'>步骤5：模型匹配过程</div>", unsafe_allow_html=True)
    st.caption("支持上传文件或输入文本，对简历与岗位进行匹配评估")
    c1, c2 = st.columns(2)
    with c1:
        rtext = st.text_area("简历文本", height=140, key="admin_match_resume")
        rfile = st.file_uploader("上传简历文件（txt/md/docx/pdf）", type=["txt", "md", "docx", "pdf"], key="admin_match_resume_file")
    with c2:
        jtext = st.text_area("岗位文本", height=140, key="admin_match_job")
        jfile = st.file_uploader("上传岗位文件（txt/md/docx/pdf）", type=["txt", "md", "docx", "pdf"], key="admin_match_job_file")
    if st.button("执行匹配", key="admin_do_match"):
        if rfile is not None and not rtext.strip():
            rtext = _read_uploaded_text(rfile)
        if jfile is not None and not jtext.strip():
            jtext = _read_uploaded_text(jfile)
        if not rtext.strip() or not jtext.strip():
            st.warning("请填写简历与岗位文本")
        else:
            res, err = _post_with_progress(base, "/match", {"resume_text": rtext, "job_text": jtext}, 8.0, "匹配评估")
            if err:
                st.error(err)
            else:
                _render_match_result(res)

    st.markdown("<div class='step-title'>步骤6：模型训练过程与评估结果</div>", unsafe_allow_html=True)
    st.caption("从日志读取训练/评估曲线，详见‘流程监控’页签")
    st.link_button("打开流程监控", "#")

    st.markdown("<div class='step-title'>步骤7：预测测试（手动上传）</div>", unsafe_allow_html=True)
    st.caption("选择实体预测或匹配评分，上传测试文本进行验证")
    mode = st.radio("测试类型", ["NER实体预测", "匹配评分"], index=0, key="admin_test_mode")
    tf = st.file_uploader("上传测试文件", type=["txt", "md", "docx", "pdf"], key="admin_test_file")
    tt = st.text_area("或粘贴文本", height=140, key="admin_test_text")
    if st.button("执行测试", key="admin_do_test"):
        ttext = tt
        if tf is not None and not ttext.strip():
            ttext = _read_uploaded_text(tf)
        if not ttext.strip():
            st.warning("请提供测试文本")
        else:
            if mode == "NER实体预测":
                res, err = api_post(base, "/predict", {"text": ttext}, timeout=6.0)
                if err:
                    st.error(err)
                else:
                    st.json(res)
            else:
                jt = st.text_area("岗位文本（匹配）", height=120, key="admin_test_job_text")
                if not jt.strip():
                    st.warning("请填写岗位文本")
                else:
                    res, err = api_post(base, "/match", {"resume_text": ttext, "job_text": jt}, timeout=8.0)
                    if err:
                        st.error(err)
                    else:
                        _render_match_result(res)

if __name__ == "__main__":
    main()